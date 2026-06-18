# -*- coding: utf-8 -*-
"""
Tao cac so do ky thuat cho bao cao chatbot RAG va chen vao Word.
"""

import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

import os, shutil
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
from matplotlib.patches import FancyArrowPatch, FancyBboxPatch
import matplotlib.patheffects as pe
import numpy as np
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

IMG_DIR = "diagram_imgs"
os.makedirs(IMG_DIR, exist_ok=True)

SRC = "Bao_Cao_Do_An_UIT_Chatbot_FINAL.docx"
DST = "Bao_Cao_Do_An_UIT_Chatbot_WITH_FIGS.docx"
shutil.copy2(SRC, DST)

# ── Helpers ──────────────────────────────────────────────────────────────────

BLUE   = "#2563EB"
GREEN  = "#16A34A"
ORANGE = "#EA580C"
GRAY   = "#6B7280"
LIGHT  = "#EFF6FF"
RED    = "#DC2626"

def box(ax, x, y, w, h, label, color=BLUE, fontsize=9, text_color="white", style="round,pad=0.1"):
    ax.add_patch(FancyBboxPatch((x-w/2, y-h/2), w, h,
        boxstyle=style, facecolor=color, edgecolor="white", linewidth=1.5))
    ax.text(x, y, label, ha="center", va="center", fontsize=fontsize,
            color=text_color, fontweight="bold", wrap=True,
            multialignment="center")

def arrow(ax, x1, y1, x2, y2, color="#374151", label=""):
    ax.annotate("", xy=(x2, y2), xytext=(x1, y1),
        arrowprops=dict(arrowstyle="->", color=color, lw=1.5))
    if label:
        mx, my = (x1+x2)/2, (y1+y2)/2
        ax.text(mx+0.05, my, label, fontsize=7, color=color, style="italic")

def save(fig, name):
    path = os.path.join(IMG_DIR, name)
    fig.savefig(path, dpi=150, bbox_inches="tight", facecolor="white")
    plt.close(fig)
    print(f"  [saved] {path}")
    return path

# ── Hình 3.1: Quy trình thực hiện đề tài ─────────────────────────────────────
def fig_31():
    fig, ax = plt.subplots(figsize=(10, 3.5))
    ax.set_xlim(0, 10); ax.set_ylim(0, 3.5); ax.axis("off")
    steps = [
        ("1. Thu thập &\nXử lý dữ liệu", 1.0, BLUE),
        ("2. Xây dựng\nCSDL Vector", 3.0, GREEN),
        ("3. Xây dựng\nHệ thống RAG", 5.0, ORANGE),
        ("4. Thực nghiệm\n& Đánh giá", 7.0, "#7C3AED"),
        ("5. Cài đặt\nDemo", 9.0, RED),
    ]
    y = 1.75
    for label, x, color in steps:
        box(ax, x, y, 1.7, 1.0, label, color=color, fontsize=8.5)
    for i in range(len(steps)-1):
        x1 = steps[i][1] + 0.85
        x2 = steps[i+1][1] - 0.85
        arrow(ax, x1, y, x2, y)
    ax.set_title("Hình 3.1: Quy trình thực hiện đề tài", fontsize=11, fontweight="bold", pad=8)
    return save(fig, "h31_quy_trinh.png")

# ── Hình 3.2: Kiến trúc tổng thể hệ thống ────────────────────────────────────
def fig_32():
    fig, ax = plt.subplots(figsize=(10, 6))
    ax.set_xlim(0, 10); ax.set_ylim(0, 6); ax.axis("off")

    layers = [
        (5, 5.4, 9, 0.8, "TẦNG GIAO DIỆN\nStreamlit Web UI", BLUE),
        (5, 4.3, 9, 0.8, "TẦNG DỊCH VỤ API\nFastAPI  /chat  /history  /health", GREEN),
        (5, 3.2, 9, 0.8, "TẦNG LÕI RAG\nTiền xử lý → Truy xuất → Sàng lọc → Tổng hợp", ORANGE),
        (5, 2.1, 9, 0.8, "TẦNG LƯU TRỮ\nQdrant (Vector DB)   PostgreSQL (History)", "#7C3AED"),
        (5, 1.0, 9, 0.8, "TẦNG HẠ TẦNG\nDocker Compose   BGE-M3 (Embed)   Gemini API (LLM)", GRAY),
    ]
    for x, y, w, h, label, color in layers:
        box(ax, x, y, w, h, label, color=color, fontsize=9)

    for i in range(len(layers)-1):
        y1 = layers[i][1] - layers[i][3]/2
        y2 = layers[i+1][1] + layers[i+1][3]/2
        ax.annotate("", xy=(5, y2), xytext=(5, y1),
            arrowprops=dict(arrowstyle="<->", color="#374151", lw=1.2))

    ax.set_title("Hình 3.2: Kiến trúc tổng thể hệ thống chatbot RAG", fontsize=11, fontweight="bold", pad=8)
    return save(fig, "h32_kien_truc.png")

# ── Hình 3.3: Luồng xử lý một truy vấn ──────────────────────────────────────
def fig_33():
    fig, ax = plt.subplots(figsize=(12, 3.8))
    ax.set_xlim(0, 12); ax.set_ylim(0, 3.8); ax.axis("off")

    nodes = [
        (1.0,  2.0, "Người\ndùng", "#1D4ED8"),
        (2.8,  2.0, "Tiền xử lý\ntruy vấn", BLUE),
        (4.8,  2.0, "Hybrid\nRetrieval", GREEN),
        (6.8,  2.0, "Reranking\n(BGE-reranker)", ORANGE),
        (8.8,  2.0, "Tổng hợp\n(Gemini LLM)", "#7C3AED"),
        (10.8, 2.0, "Câu trả\nlời", RED),
    ]
    for x, y, label, color in nodes:
        box(ax, x, y, 1.6, 0.9, label, color=color, fontsize=8)

    for i in range(len(nodes)-1):
        arrow(ax, nodes[i][0]+0.8, nodes[i][1], nodes[i+1][0]-0.8, nodes[i+1][1])

    # Chitchat branch
    box(ax, 2.8, 0.6, 1.6, 0.6, "Chitchat\n→ trả lời trực tiếp", color=GRAY, fontsize=7)
    ax.annotate("", xy=(2.8, 0.9), xytext=(2.8, 1.55),
        arrowprops=dict(arrowstyle="->", color=GRAY, lw=1.2))
    ax.text(2.2, 1.2, "chitchat?", fontsize=7, color=GRAY, style="italic")

    # Qdrant
    box(ax, 4.8, 0.6, 1.6, 0.6, "Qdrant\n(Vector Store)", color="#0891B2", fontsize=7)
    ax.annotate("", xy=(4.8, 1.55), xytext=(4.8, 0.9),
        arrowprops=dict(arrowstyle="<->", color="#0891B2", lw=1.2))

    ax.set_title("Hình 3.3: Luồng xử lý một truy vấn trong hệ thống", fontsize=11, fontweight="bold", pad=8)
    return save(fig, "h33_luong_xu_ly.png")

# ── Hình 3.4: Pipeline thu thập và xử lý dữ liệu ────────────────────────────
def fig_34():
    fig, ax = plt.subplots(figsize=(11, 4.5))
    ax.set_xlim(0, 11); ax.set_ylim(0, 4.5); ax.axis("off")

    # Sources
    sources = [
        (1, 3.8, "PDF\nVăn bản", BLUE),
        (1, 2.5, "PDF\nScan", ORANGE),
        (1, 1.2, "DOCX\n& JSON", GREEN),
    ]
    for x, y, label, color in sources:
        box(ax, x, y, 1.4, 0.8, label, color=color, fontsize=8)

    # Extract
    box(ax, 3.2, 3.8, 1.4, 0.8, "PyMuPDF\n(direct)", BLUE, fontsize=8)
    box(ax, 3.2, 2.5, 1.4, 0.8, "OCR\n(Tesseract)", ORANGE, fontsize=8)
    box(ax, 3.2, 1.2, 1.4, 0.8, "python-docx\nparser", GREEN, fontsize=8)

    for i, (sx, sy, _, _) in enumerate(sources):
        ex = [3.2, 3.2, 3.2][i]
        ey = [3.8, 2.5, 1.2][i]
        arrow(ax, sx+0.7, sy, ex-0.7, ey)

    # Clean
    box(ax, 5.5, 2.5, 1.5, 0.8, "Làm sạch\n& Chuẩn hóa", "#7C3AED", fontsize=8)
    for ey in [3.8, 2.5, 1.2]:
        ax.annotate("", xy=(5.5-0.2, 2.5), xytext=(3.2+0.7, ey),
            arrowprops=dict(arrowstyle="->", color="#374151", lw=1, connectionstyle="arc3,rad=0.0"))

    # Metadata
    box(ax, 7.5, 2.5, 1.5, 0.8, "Gắn\nMetadata", "#0891B2", fontsize=8)
    arrow(ax, 5.5+0.75, 2.5, 7.5-0.75, 2.5)

    # Store
    box(ax, 9.5, 2.5, 1.5, 0.8, "Kho văn bản\nJSON/DB", "#16A34A", fontsize=8)
    arrow(ax, 7.5+0.75, 2.5, 9.5-0.75, 2.5)

    ax.text(0.2, 4.3, "Nguồn", fontsize=8, color=GRAY, style="italic")
    ax.text(2.5, 4.3, "Trích xuất", fontsize=8, color=GRAY, style="italic")
    ax.text(5.1, 4.3, "Làm sạch", fontsize=8, color=GRAY, style="italic")
    ax.text(7.1, 4.3, "Metadata", fontsize=8, color=GRAY, style="italic")
    ax.text(9.1, 4.3, "Lưu trữ", fontsize=8, color=GRAY, style="italic")

    ax.set_title("Hình 3.4: Pipeline thu thập và xử lý dữ liệu", fontsize=11, fontweight="bold", pad=8)
    return save(fig, "h34_pipeline_data.png")

# ── Hình 3.5: Quy trình xây dựng CSDL vector ─────────────────────────────────
def fig_35():
    fig, ax = plt.subplots(figsize=(10, 3.2))
    ax.set_xlim(0, 10); ax.set_ylim(0, 3.2); ax.axis("off")

    steps = [
        (1.2, 1.6, "Văn bản\nthô", BLUE),
        (3.0, 1.6, "Chunking\n(512 từ,\noverlap 64)", GREEN),
        (5.0, 1.6, "BGE-M3\nEmbedding\n(1024-dim)", ORANGE),
        (7.0, 1.6, "Qdrant\nIndexing\n(cosine)", "#7C3AED"),
        (9.0, 1.6, "Collection\ndocuments", RED),
    ]
    for x, y, label, color in steps:
        box(ax, x, y, 1.55, 1.1, label, color=color, fontsize=8)
    for i in range(len(steps)-1):
        arrow(ax, steps[i][0]+0.78, steps[i][1], steps[i+1][0]-0.78, steps[i+1][1])

    ax.set_title("Hình 3.5: Quy trình xây dựng cơ sở dữ liệu vector", fontsize=11, fontweight="bold", pad=8)
    return save(fig, "h35_csdl_vector.png")

# ── Hình 5.1: Sơ đồ thiết kế API ─────────────────────────────────────────────
def fig_51():
    fig, ax = plt.subplots(figsize=(10, 5))
    ax.set_xlim(0, 10); ax.set_ylim(0, 5); ax.axis("off")

    # Client
    box(ax, 1.2, 2.5, 1.6, 0.8, "Streamlit\nClient", BLUE, fontsize=8)

    # Endpoints
    eps = [
        (4.0, 4.2, "POST /chat", GREEN),
        (4.0, 3.1, "GET /history/{id}", ORANGE),
        (4.0, 2.0, "DELETE /history/{id}", RED),
        (4.0, 0.9, "GET /health", GRAY),
    ]
    for x, y, label, color in eps:
        box(ax, x, y, 2.0, 0.7, label, color=color, fontsize=8)
        arrow(ax, 1.2+0.8, 2.5, x-1.0, y)

    # Backend services
    box(ax, 7.5, 3.6, 2.0, 0.8, "RAG Pipeline\n(Chương 3)", ORANGE, fontsize=8)
    box(ax, 7.5, 2.2, 2.0, 0.8, "PostgreSQL\n(History)", "#7C3AED", fontsize=8)
    box(ax, 7.5, 0.9, 2.0, 0.8, "Health Check", GRAY, fontsize=8)

    arrow(ax, 4.0+1.0, 4.2, 7.5-1.0, 3.6)
    arrow(ax, 4.0+1.0, 3.1, 7.5-1.0, 2.2)
    arrow(ax, 4.0+1.0, 2.0, 7.5-1.0, 2.2)
    arrow(ax, 4.0+1.0, 0.9, 7.5-1.0, 0.9)

    ax.text(3.8, 4.75, "FastAPI Endpoints", fontsize=9, color=GRAY, style="italic", fontweight="bold")
    ax.set_title("Hình 5.1: Sơ đồ thiết kế API và luồng dữ liệu", fontsize=11, fontweight="bold", pad=8)
    return save(fig, "h51_api_design.png")

# ── Generate tất cả hình ─────────────────────────────────────────────────────
print("Dang tao cac so do...")
paths = {
    "Hình 3.1": fig_31(),
    "Hình 3.2": fig_32(),
    "Hình 3.3": fig_33(),
    "Hình 3.4": fig_34(),
    "Hình 3.5": fig_35(),
    "Hình 5.1": fig_51(),
}
print(f"Da tao {len(paths)} so do\n")

# ── Chèn ảnh vào Word ─────────────────────────────────────────────────────────
doc = Document(DST)

placeholder_map = {
    "[Hình 3.1: Quy trình thực hiện đề tài — sẽ bổ sung]":           paths["Hình 3.1"],
    "[Hình 3.2: Kiến trúc tổng thể hệ thống chatbot RAG — sẽ bổ sung]": paths["Hình 3.2"],
    "[Hình 3.3: Luồng xử lý một truy vấn trong hệ thống — sẽ bổ sung]": paths["Hình 3.3"],
    "[Hình 3.4: Pipeline thu thập và xử lý dữ liệu — sẽ bổ sung]":   paths["Hình 3.4"],
    "[Hình 3.5: Quy trình xây dựng cơ sở dữ liệu vector — sẽ bổ sung]": paths["Hình 3.5"],
    "[Hình 5.1: Sơ đồ thiết kế API và luồng dữ liệu — sẽ bổ sung]":  paths["Hình 5.1"],
}

ui_placeholders = {
    "[Hình 5.2: Giao diện trò chuyện của hệ thống — sẽ bổ sung]",
    "[Hình 5.3: Minh họa tương tác hỏi đáp cơ bản — sẽ bổ sung]",
    "[Hình 5.4: Minh họa hội thoại nhiều lượt với viết lại truy vấn — sẽ bổ sung]",
    "[Hình 5.5: Minh họa câu trả lời tổng hợp từ nhiều nguồn — sẽ bổ sung]",
    "[Hình 5.6: Minh họa câu trả lời dạng bảng so sánh — sẽ bổ sung]",
    "[Hình 5.7: Minh họa từ chối câu hỏi ngoài phạm vi — sẽ bổ sung]",
}

def clear_para_and_add_image(para, img_path, caption_text):
    """Xoa text placeholder, chen anh va caption."""
    # Xoa tat ca run cu
    p_elem = para._element
    for r in list(p_elem.findall(qn("w:r"))):
        p_elem.remove(r)
    # Chen anh vao para nay
    run = para.add_run()
    run.add_picture(img_path, width=Inches(5.5))
    para.alignment = 1  # WD_ALIGN_PARAGRAPH.CENTER
    # Them caption moi sau para
    from docx.oxml import OxmlElement as OE
    cap_p = OE("w:p")
    cap_r = OE("w:r")
    cap_rPr = OE("w:rPr")
    cap_i = OE("w:i"); cap_rPr.append(cap_i)
    cap_r.append(cap_rPr)
    cap_t = OE("w:t")
    cap_t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    cap_t.text = caption_text
    cap_r.append(cap_t)
    cap_p.append(cap_r)
    p_elem.addnext(cap_p)

for para in doc.paragraphs:
    txt = para.text.strip()
    if txt in placeholder_map:
        img_path = placeholder_map[txt]
        # Extract caption tu text (bo dau [...])
        caption = txt.replace("[", "").replace("]", "").replace(" — sẽ bổ sung", "")
        clear_para_and_add_image(para, img_path, caption)
        print(f"[OK] Da chen: {caption}")
    elif txt in ui_placeholders:
        # Thay bang text do nhac nho screenshot
        p_elem = para._element
        for r in list(p_elem.findall(qn("w:r"))):
            p_elem.remove(r)
        r_new = OxmlElement("w:r")
        rPr = OxmlElement("w:rPr")
        col = OxmlElement("w:color"); col.set(qn("w:val"), "FF0000")
        rPr.append(col)
        i_elem = OxmlElement("w:i"); rPr.append(i_elem)
        r_new.append(rPr)
        t = OxmlElement("w:t")
        t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
        caption = txt.replace("[", "").replace("]", "").replace(" — sẽ bổ sung", "")
        t.text = f"[CẦN CHỤP MÀN HÌNH THỰC TẾ] {caption}"
        r_new.append(t)
        p_elem.append(r_new)
        print(f"[NOTE] Placeholder UI: {caption}")

doc.save(DST)
print(f"\nDa luu: {DST}")
