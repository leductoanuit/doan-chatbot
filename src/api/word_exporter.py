"""Word (.docx) report generation — Type A (chat export) and Type B (technical report)."""

import tempfile
from datetime import datetime
from typing import Dict, List

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import RGBColor


# UIT brand colour (blue)
_UIT_BLUE = RGBColor(0x1A, 0x56, 0xDB)


def _add_centered(doc: Document, text: str, bold: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def export_chat_report(history: List[Dict]) -> str:
    """Type A: Export conversation history + sources as .docx.

    Args:
        history: List of role/content dicts from the chat session.

    Returns:
        Path to the generated temporary .docx file.
    """
    doc = Document()

    # Title block
    title = doc.add_heading("Báo Cáo Tư Vấn Đào Tạo", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_centered(doc, "Trường Đại học Công nghệ Thông tin - ĐHQG TP.HCM")
    doc.add_paragraph(f"Ngày xuất: {datetime.now().strftime('%d/%m/%Y %H:%M')}")
    user_turns = [h for h in history if h.get("role") == "user"]
    doc.add_paragraph(f"Số lượng câu hỏi: {len(user_turns)}")
    doc.add_paragraph()

    # Conversation
    doc.add_heading("Nội Dung Trao Đổi", level=2)
    for msg in history:
        role = msg.get("role", "")
        content = msg.get("content", "")
        if role == "system":
            continue

        p = doc.add_paragraph()
        if role == "user":
            run = p.add_run("Sinh viên: ")
            run.bold = True
            p.add_run(content)
        elif role == "assistant":
            run = p.add_run("Tư vấn viên: ")
            run.bold = True
            run.font.color.rgb = _UIT_BLUE
            p.add_run(content)
            doc.add_paragraph()  # Spacing after each assistant turn

    # Footer
    doc.add_page_break()
    doc.add_heading("Ghi Chú", level=2)
    doc.add_paragraph(
        "Các thông tin trên được cung cấp từ hệ thống chatbot tư vấn đào tạo UIT. "
        "Để biết thêm chi tiết, vui lòng liên hệ phòng đào tạo hoặc "
        "truy cập website daa.uit.edu.vn."
    )

    return _save_temp(doc)


def export_technical_report(
    history: List[Dict],
    system_metrics: Dict = None,
    eval_results: Dict = None,
) -> str:
    """Type B: Export technical report for thesis committee (.docx).

    Includes system architecture summary, evaluation metrics, and conversation stats.

    Args:
        history: Chat history for statistics.
        system_metrics: Optional dict with performance numbers.
        eval_results: Optional dict with evaluation scores.

    Returns:
        Path to the generated temporary .docx file.
    """
    doc = Document()
    now = datetime.now().strftime("%d/%m/%Y %H:%M")

    # Cover
    title = doc.add_heading("Báo Cáo Kỹ Thuật Hệ Thống Chatbot Tư Vấn Đào Tạo", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_centered(doc, "Trường Đại học Công nghệ Thông tin - ĐHQG TP.HCM")
    _add_centered(doc, f"Ngày xuất: {now}")
    doc.add_page_break()

    # 1. System Architecture
    doc.add_heading("1. Kiến Trúc Hệ Thống", level=2)
    arch_rows = [
        ("Thành phần", "Công nghệ"),
        ("Backend API", "Python 3.11 + FastAPI"),
        ("Mô hình ngôn ngữ", "Gemini 2.0 Flash (Google AI)"),
        ("Vector Database", "MongoDB Atlas Free Tier (768-dim cosine)"),
        ("Embedding Model", "dangvantuan/vietnamese-embedding (PhoBERT)"),
        ("RAG Framework", "LlamaIndex + Hybrid Search"),
        ("LLM Serving", "Gemini API"),
        ("OCR", "DeepSeek VL2 (Google Colab)"),
        ("Frontend", "Streamlit"),
    ]
    table = doc.add_table(rows=len(arch_rows), cols=2)
    table.style = "Table Grid"
    for i, (col1, col2) in enumerate(arch_rows):
        row = table.rows[i]
        row.cells[0].text = col1
        row.cells[1].text = col2
        if i == 0:
            for cell in row.cells:
                for run in cell.paragraphs[0].runs:
                    run.bold = True

    doc.add_paragraph()

    # 2. Performance Metrics
    doc.add_heading("2. Kết Quả Đánh Giá Hiệu Suất", level=2)
    if system_metrics:
        metrics_rows = [("Chỉ số", "Giá trị")] + list(system_metrics.items())
        m_table = doc.add_table(rows=len(metrics_rows), cols=2)
        m_table.style = "Table Grid"
        for i, (k, v) in enumerate(metrics_rows):
            m_table.rows[i].cells[0].text = str(k)
            m_table.rows[i].cells[1].text = str(v)
    else:
        doc.add_paragraph(
            "Chưa có dữ liệu đánh giá. Chạy tests/eval-e2e.py để thu thập kết quả."
        )

    doc.add_paragraph()

    # 3. Evaluation Results
    doc.add_heading("3. Kết Quả Đánh Giá Chất Lượng", level=2)
    if eval_results:
        for metric, value in eval_results.items():
            doc.add_paragraph(f"{metric}: {value}", style="List Bullet")
    else:
        doc.add_paragraph("Chưa có kết quả đánh giá chất lượng.")

    doc.add_paragraph()

    # 4. Session Statistics
    doc.add_heading("4. Thống Kê Phiên Tư Vấn", level=2)
    user_msgs = [h for h in history if h.get("role") == "user"]
    doc.add_paragraph(f"Tổng số câu hỏi: {len(user_msgs)}")
    doc.add_paragraph(f"Thời gian: {now}")

    doc.add_page_break()
    doc.add_heading("5. Ghi Chú", level=2)
    doc.add_paragraph(
        "Báo cáo này được tạo tự động bởi hệ thống chatbot tư vấn đào tạo UIT. "
        "Đây là đồ án tốt nghiệp — mọi số liệu chỉ mang tính tham khảo."
    )

    return _save_temp(doc, prefix="bao-cao-ky-thuat")


def _save_temp(doc: Document, prefix: str = "bao-cao-tu-van") -> str:
    """Save Document to a named temp file and return its path."""
    tmp = tempfile.NamedTemporaryFile(
        delete=False,
        suffix=".docx",
        prefix=f"{prefix}-",
    )
    doc.save(tmp.name)
    tmp.close()
    return tmp.name
