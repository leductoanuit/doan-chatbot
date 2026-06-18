# -*- coding: utf-8 -*-
import sys, io, shutil
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SRC = "Bao_Cao_Do_An_UIT_Chatbot_WITH_FIGS.docx"
TMP = "tmp_fix.docx"
shutil.copy2(SRC, TMP)
doc = Document(TMP)

# ── Buoc 1: Xoa cac para da chen sai ─────────────────────────────────────────
markers = ["Bước 1", "Bước 2", "Bước 3", "Bước 4", "Bước 5",
           "Hệ thống được cài đặt hoàn toàn cục bộ"]

to_remove = []
for p in doc.paragraphs:
    if any(p.text.startswith(m) for m in markers):
        to_remove.append(p)
    if p.text.strip() == "Quy trình cài đặt" and "Heading" in p.style.name:
        to_remove.append(p)

for p in to_remove:
    p._element.getparent().remove(p._element)
    print(f"[DEL] {p.text[:60]}")


# ── Buoc 2: Helper tao para mau do ───────────────────────────────────────────
def make_red(text, style_id=None):
    p = OxmlElement("w:p")
    if style_id:
        pPr = OxmlElement("w:pPr")
        ps = OxmlElement("w:pStyle"); ps.set(qn("w:val"), style_id)
        pPr.append(ps); p.append(pPr)
    r = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    col = OxmlElement("w:color"); col.set(qn("w:val"), "FF0000")
    rPr.append(col); r.append(rPr)
    t = OxmlElement("w:t")
    t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    t.text = text; r.append(t); p.append(r)
    return p


# ── Buoc 3: Tim diem chen (truoc "Thiet ke API") ─────────────────────────────
ref = None
for p in doc.paragraphs:
    if p.text.strip() == "Thiết kế API":
        ref = p; break

assert ref, "Khong tim thay 'Thiet ke API'"

# ── Buoc 4: Tao noi dung theo thu tu dung, chen bang addprevious nguoc ───────
# Thu tu hien thi mong muon: Heading → intro → B1 → B2 → B3 → B4 → B5
# Chen bang addprevious theo thu tu NGUOC (B5 truoc, Heading cuoi)
# ⟹ B5 chen vao → nam ngay truoc Thiet ke API
#    B4 chen vao → nam truoc B5
#    ... Heading chen cuoi → nam dau tien

contents = [
    ("h2",   "Quy trình cài đặt"),
    ("para", "Hệ thống được cài đặt hoàn toàn cục bộ qua Docker Compose. "
             "Các bước thực hiện như sau:"),
    ("list", "Bước 1 — Yêu cầu hệ thống: Python 3.12+, Docker Desktop (Windows/macOS) "
             "hoặc Docker Engine (Linux), RAM tối thiểu 8 GB (khuyến nghị 16 GB), "
             "ổ cứng trống tối thiểu 10 GB (cho model BGE-M3 ~1.5 GB và dữ liệu vector). "
             "Không yêu cầu GPU — reranker chạy trên CPU."),
    ("list", "Bước 2 — Cài đặt hạ tầng bằng Docker Compose: đảm bảo Docker Desktop đang chạy, "
             "sau đó chạy lệnh docker-compose up -d qdrant postgres "
             "để khởi động Qdrant (cổng 6333) và PostgreSQL 16 (cổng 5432) ở chế độ nền."),
    ("list", "Bước 3 — Cấu hình biến môi trường: sao chép file .env.example thành .env "
             "và điền các giá trị: GEMINI_API_KEY (lấy từ Google AI Studio), "
             "QDRANT_HOST=localhost, QDRANT_PORT=6333, "
             "POSTGRES_URL=postgresql://user:pass@localhost:5432/chatbot. "
             "Cài đặt thư viện Python: pip install -r requirements.txt."),
    ("list", "Bước 4 — Nạp dữ liệu và khởi chạy ứng dụng: chạy python ingest.py "
             "để nạp dữ liệu vào Qdrant vector store, sau đó chạy "
             "uvicorn main:app --reload (FastAPI, cổng 8000) "
             "và streamlit run app.py (giao diện, cổng 8501) trên hai terminal riêng."),
    ("list", "Bước 5 — Truy cập giao diện: mở trình duyệt và truy cập "
             "http://localhost:8501 để sử dụng giao diện Streamlit. "
             "Nhập câu hỏi bằng tiếng Việt vào khung chat và nhấn Enter để nhận câu trả lời."),
]

# Chen theo thu tu xuoi - moi elem addprevious truoc Thiet ke API
# → heading vao truoc, B5 vao sau cung, ket qua: heading→intro→B1→B2→B3→B4→B5
for kind, text in contents:
    style_id = "Heading2" if kind == "h2" else ("ListParagraph" if kind == "list" else None)
    elem = make_red(text, style_id)
    ref._element.addprevious(elem)
    print(f"[INS] {text[:60]}")

doc.save(TMP)
shutil.move(TMP, SRC)
print(f"\n[OK] Da luu: {SRC}")
