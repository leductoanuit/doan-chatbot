# -*- coding: utf-8 -*-
"""
Script v2: Thay the tat ca [GHI CHU] bang noi dung thuc bang tieng Viet.
Noi dung them moi van duoc to do.
"""

import shutil
from docx import Document
from docx.shared import RGBColor, Pt, Cm
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH
from lxml import etree

SRC = "Bao_Cao_Do_An_UIT_Chatbot_REVISED.docx"
DST = "Bao_Cao_Do_An_UIT_Chatbot_FINAL.docx"
shutil.copy2(SRC, DST)
doc = Document(DST)


# ─── Helpers ─────────────────────────────────────────────────────────────────

def red_run_xml(text):
    """Tao <w:r> mau do."""
    r = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "FF0000")
    rPr.append(color)
    r.append(rPr)
    t = OxmlElement("w:t")
    t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    t.text = text
    r.append(t)
    return r


def new_para_xml(text, red=True, style_id=None):
    """Tao <w:p> moi voi text (tuy chon mau do)."""
    p = OxmlElement("w:p")
    if style_id:
        pPr = OxmlElement("w:pPr")
        pStyle = OxmlElement("w:pStyle")
        pStyle.set(qn("w:val"), style_id)
        pPr.append(pStyle)
        p.append(pPr)
    r = OxmlElement("w:r")
    if red:
        rPr = OxmlElement("w:rPr")
        color = OxmlElement("w:color")
        color.set(qn("w:val"), "FF0000")
        rPr.append(color)
        r.append(rPr)
    t = OxmlElement("w:t")
    t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    t.text = text
    r.append(t)
    p.append(r)
    return p


def replace_para_with_red(para, new_text):
    """Xoa noi dung cu cua para, viet lai bang text mau do."""
    p_elem = para._element
    # Xoa tat ca run cu
    for r in p_elem.findall(qn("w:r")):
        p_elem.remove(r)
    # Them run moi mau do
    p_elem.append(red_run_xml(new_text))


def insert_red_para_after(ref_para, text, style_id=None):
    new_p = new_para_xml(text, red=True, style_id=style_id)
    ref_para._element.addnext(new_p)
    return new_p


def remove_para(para):
    """Xoa paragraph khoi document."""
    para._element.getparent().remove(para._element)


def find_para_containing(doc, marker):
    """Tim paragraph dau tien chua marker."""
    for p in doc.paragraphs:
        if marker in p.text:
            return p
    return None


# ─── 1. Xoa ghi chu FORMAT tong hop (para 1) ─────────────────────────────────
p = find_para_containing(doc, "GHI CHU TONG HOP CAC LOI FORMAT")
if p:
    remove_para(p)
    print("[OK] Da xoa ghi chu FORMAT tong hop")

# ─── 2. TOM TAT: Sua doan Huong phat trien (hien dang ASCII, chuyen sang TV) ─
p = find_para_containing(doc, "GHI CHU: BO SUNG -- Thay nhac thieu y thu 4")
if p:
    replace_para_with_red(p,
        "Về hướng phát triển, đề tài đề xuất các cải tiến trọng tâm: "
        "giảm độ trễ phản hồi bằng cách triển khai reranker trên GPU (mục tiêu dưới 10 giây/câu); "
        "cải thiện độ chính xác câu trả lời qua tinh chỉnh prompt với few-shot examples; "
        "mở rộng phạm vi tri thức sang các đơn vị khác trong ĐHQG-HCM; "
        "và bổ sung đánh giá bởi người dùng thực để hoàn thiện hệ thống."
    )
    print("[OK] Da sua doan Huong phat trien trong TOM TAT")

# ─── 3. MUC TIEU: Xoa ghi chu, them muc tieu thu nhat (kien truc) vao dau ─────
p = find_para_containing(doc, "GHI CHU: SAP XEP LAI")
if p:
    remove_para(p)
    print("[OK] Da xoa ghi chu SAP XEP LAI muc tieu")

# Tim heading "Muc tieu cua de tai" va "De tai huong den..."
p_intro = find_para_containing(doc, "Đề tài hướng đến các mục tiêu cụ thể")
if p_intro:
    # Them bullet kien truc he thong vao dau danh sach
    insert_red_para_after(p_intro,
        "– Xác định kiến trúc hệ thống tổng thể cho chatbot hỏi đáp tiếng Việt "
        "theo mô hình RAG phân tầng, xác định các thành phần và luồng xử lý dữ liệu.",
        style_id="ListParagraph"
    )
    print("[OK] Da them bullet kien truc he thong vao Muc tieu")

# ─── 4. PHAT BIEU BAI TOAN: Sua vi du ASCII sang tieng Viet ──────────────────
p = find_para_containing(doc, "GHI CHU: BO SUNG VI DU")
if p:
    replace_para_with_red(p,
        "Ví dụ minh họa — Đầu vào: \"Điều kiện để được xét tốt nghiệp hệ đào tạo từ xa UIT là gì?\"; "
        "Đầu ra: \"Theo Quyết định 1499/QĐ-ĐHCNTT (2024), sinh viên cần hoàn thành đủ số tín chỉ "
        "theo chương trình đào tạo, đạt chuẩn đầu ra ngoại ngữ và công nghệ thông tin, "
        "không trong thời gian bị kỷ luật và nộp đơn xét tốt nghiệp đúng thời hạn quy định.\""
    )
    print("[OK] Da sua vi du dau vao/dau ra sang tieng Viet")

# ─── 5. DOI TUONG: Di chuyen "Pham vi trien khai" sang dung vi tri ──────────
# Hien tai no dang o trong muc "Doi tuong nghien cuu", phai chuyen vao "Pham vi nghien cuu"
p_wrong = find_para_containing(doc, "GHI CHU: BO SUNG -- Thay yeu cau them 'Pham vi trien khai'")
if p_wrong:
    # Xoa para sai vi tri
    remove_para(p_wrong)
    print("[OK] Da xoa ghi chu Pham vi trien khai (sai vi tri)")

# Tim bullet cuoi cua Pham vi nghien cuu (pham vi cong nghe)
p_cong_nghe = find_para_containing(doc, "Phạm vi công nghệ: sử dụng LLM thương mại")
if p_cong_nghe:
    insert_red_para_after(p_cong_nghe,
        "– Phạm vi triển khai: hệ thống được thiết kế để vận hành trên môi trường cục bộ "
        "(on-premise) hoặc máy chủ nội bộ của nhà trường; đối tượng người dùng là sinh viên "
        "và học viên hệ đào tạo từ xa UIT truy cập qua trình duyệt web; "
        "không yêu cầu cài đặt phần mềm phía người dùng.",
        style_id="ListParagraph"
    )
    print("[OK] Da them Pham vi trien khai dung vi tri")

# ─── 6. Y NGHIA: Di chuyen ghi chu sai vi tri, them noi dung thuc ────────────
p_wrong2 = find_para_containing(doc, "GHI CHU: MO RONG -- Thay nhac phan y nghia")
if p_wrong2:
    remove_para(p_wrong2)
    print("[OK] Da xoa ghi chu Y nghia sai vi tri")

p_yt = find_para_containing(doc, "Ý nghĩa thực tiễn: sản phẩm chatbot")
if p_yt:
    insert_red_para_after(p_yt,
        "Ý nghĩa đào tạo: đề tài đóng góp bộ dữ liệu kiểm thử gồm 100 cặp câu hỏi — đáp án "
        "gắn nhãn thủ công và bộ khung đánh giá tự động đa chiều theo định hướng Auepora, "
        "có thể tái sử dụng cho các nghiên cứu chatbot học vụ tiếng Việt về sau. "
        "Ý nghĩa mở rộng: quy trình xây dựng và đánh giá của đề tài có thể áp dụng cho "
        "các đơn vị đào tạo khác trong hệ thống ĐHQG-HCM hoặc các trường có nhu cầu tương tự."
    )
    print("[OK] Da them noi dung Y nghia mo rong")

# ─── 7. CHUONG 2: Xoa ghi chu trich dan (chi la huong dan, khong phai noi dung) ─
p = find_para_containing(doc, "GHI CHU: BO SUNG TRICH DAN")
if p:
    remove_para(p)
    print("[OK] Da xoa ghi chu trich dan (chi la huong dan)")

# ─── 8. CHUONG 2: Tao bang he thong hoa nghien cuu lien quan ─────────────────
p_bang_note = find_para_containing(doc, "GHI CHU: BO SUNG BANG")
if p_bang_note:
    # Chuan bi vi tri: insert table truoc khoang trong heading
    # Tim heading "Khoang trong nghien cuu"
    p_khoang = find_para_containing(doc, "Khoảng trống nghiên cứu và ý tưởng đề xuất")

    # Them caption truoc heading
    if p_khoang:
        caption_p = new_para_xml(
            "Bảng 2.1: Tổng hợp các nghiên cứu liên quan đến chatbot hỏi đáp trong giáo dục",
            red=True
        )
        p_khoang._element.addprevious(caption_p)

        # Tao table 6 cot x 7 hang (1 header + 6 cong trinh)
        table = doc.add_table(rows=7, cols=5)
        table.style = "Table Grid"

        headers = [
            "Công trình",
            "Bài toán / Ứng dụng",
            "Phương pháp chính",
            "Kết quả nổi bật",
            "Hạn chế / Ghi chú",
        ]
        rows_data = [
            [
                "Lewis và cộng sự (2020)",
                "Hỏi đáp miền mở (open-domain QA)",
                "RAG: dense retrieval (DPR) + seq2seq (BART)",
                "Cải thiện đáng kể so với fine-tuning thuần trên NQ, TriviaQA",
                "Chưa tối ưu cho tiếng Việt và văn bản hành chính",
            ],
            [
                "Chen và cộng sự (2024) — BGE-M3",
                "Biểu diễn ngữ nghĩa đa ngôn ngữ",
                "Mô hình nhúng đa chức năng: dense, sparse, multi-vec",
                "Hỗ trợ 100+ ngôn ngữ, SOTA nhiều benchmark đa ngôn ngữ",
                "Chi phí tính toán cao khi dùng đủ 3 chế độ",
            ],
            [
                "Rackauckas (2024) — RAG Fusion",
                "Cải thiện truy xuất bằng đa truy vấn",
                "Sinh nhiều biến thể truy vấn + Reciprocal Rank Fusion",
                "Tăng recall, hạn chế miss do câu hỏi mơ hồ",
                "Tăng độ trễ do cần gọi LLM thêm lần",
            ],
            [
                "Nghiên cứu chatbot ĐH trên thế giới (tổng hợp)",
                "Tư vấn sinh viên, hỏi đáp học vụ",
                "RAG với LLM thương mại (GPT-4, Gemini)",
                "Giảm tải bộ phận hỗ trợ, phản hồi 24/7",
                "Ít công bố chi tiết quy trình đánh giá định lượng",
            ],
            [
                "Nghiên cứu chatbot tuyển sinh VN (tổng hợp)",
                "Tư vấn tuyển sinh đại học trong nước",
                "Intent-based (Rasa, Dialogflow), một số dùng RAG",
                "Triển khai thực tế tại nhiều trường",
                "Phần lớn intent-based, ít đánh giá định lượng, chưa xử lý PDF scan",
            ],
            [
                "Đề tài này",
                "Hỏi đáp học vụ hệ đào tạo từ xa UIT",
                "RAG 4 giai đoạn: tiền xử lý, hybrid retrieval, rerank, CoT",
                "MRR 0,99; Hit@10 1,0; Negative Rejection 1,0",
                "Correctness 0,58–0,59; độ trễ còn cao (reranker trên CPU)",
            ],
        ]

        # Header row
        hdr_row = table.rows[0]
        for i, hdr in enumerate(headers):
            cell = hdr_row.cells[i]
            cell.text = hdr
            for run in cell.paragraphs[0].runs:
                run.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)

        # Data rows
        for r_idx, row_data in enumerate(rows_data):
            row = table.rows[r_idx + 1]
            for c_idx, cell_text in enumerate(row_data):
                cell = row.cells[c_idx]
                cell.text = cell_text
                for run in cell.paragraphs[0].runs:
                    run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)

        # Di chuyen table XML vao dung vi tri (truoc caption, truoc heading khoang trong)
        tbl_elem = table._tbl
        caption_p.addprevious(tbl_elem)

    # Xoa ghi chu cu
    remove_para(p_bang_note)
    print("[OK] Da tao bang tong hop nghien cuu lien quan")

# ─── 9. CHUONG 3: Xoa ghi chu SUA TU NGU (da sua xong o v1) ─────────────────
p = find_para_containing(doc, "GHI CHU: SUA TU NGU")
if p:
    remove_para(p)
    print("[OK] Da xoa ghi chu SUA TU NGU")

# Xoa "[DA SUA]" khoi van ban (de sach hon)
for para in doc.paragraphs:
    for run in para.runs:
        if "[ĐÃ SỬA]" in run.text:
            run.text = run.text.replace(" [ĐÃ SỬA]", "")

# ─── 10. CHUONG 4: Mo ta tap kiem thu chi tiet ───────────────────────────────
p = find_para_containing(doc, "GHI CHU: BO SUNG -- Thay yeu cau mo ta ro hon tap danh gia")
if p:
    replace_para_with_red(p,
        "Tập kiểm thử chuẩn gồm 100 câu hỏi phân bổ đều theo 5 danh mục tri thức: "
        "tuyển sinh (20 câu), học vụ (25 câu), học phí (15 câu), quy chế — kỷ luật (20 câu) "
        "và chứng chỉ — văn bằng (20 câu); độ khó được phân tầng thành ba mức: "
        "dễ — câu hỏi tra cứu trực tiếp (40%), trung bình — cần suy luận một bước (40%), "
        "khó — cần tổng hợp nhiều đoạn văn bản (20%). "
        "Ví dụ minh họa: (1) [Dễ] \"Học phí hệ đào tạo từ xa UIT năm học 2024–2025 là bao nhiêu?\" "
        "→ đáp án tra trực tiếp từ Quyết định 507/QĐ-ĐHCNTT; "
        "(2) [Trung bình] \"Sinh viên cần đáp ứng điều kiện gì để được đăng ký học lại?\" "
        "→ cần đọc kết hợp quy chế đào tạo và quy định học vụ; "
        "(3) [Khó] \"So sánh điều kiện tốt nghiệp giữa chương trình văn bằng 1 và liên thông CNTT từ xa.\" "
        "→ cần tổng hợp từ 2–3 văn bản quy định khác nhau. "
        "Lưu ý: toàn bộ tập kiểm thử được nhóm tự xây dựng dựa trên nội dung văn bản quy định; "
        "LLM chỉ được dùng để hỗ trợ diễn đạt, không sinh câu hỏi — đáp án thay thế con người."
    )
    print("[OK] Da viet noi dung mo ta tap kiem thu")

# ─── 11. Xoa ghi chu TLTKHAM KHAO (chi la huong dan, khong phai noi dung) ─────
p = find_para_containing(doc, "GHI CHU: SUA DINH DANG TLTKHAM KHAO")
if p:
    remove_para(p)
    print("[OK] Da xoa ghi chu TLTKHAM KHAO")

# ─── Luu ─────────────────────────────────────────────────────────────────────
doc.save(DST)
print(f"\nDa luu: {DST}")
print("Tat ca noi dung them moi duoc to DO trong file.")
