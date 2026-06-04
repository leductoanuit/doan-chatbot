# Build Bao_Cao_Do_An_UIT_Chatbot.docx from markdown drafts,
# inheriting ALL styles (fonts, sizes, heading numbering, margins) from the
# template file Bao_Cao_Do_An_Chatbot_2105.docx.
# Usage: python build-docx-from-drafts.py
import copy
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn

PLAN_DIR = Path(__file__).parent
DRAFTS = sorted((PLAN_DIR / "drafts").glob("chuong-*.md"))
TEMPLATE = Path(r"D:\doan-chatbot\Bao_Cao_Do_An_Chatbot_2105.docx")
OUT = Path(r"D:\doan-chatbot\Bao_Cao_Do_An_UIT_Chatbot.docx")

# Load template (styles, numbering, margins carry over), then wipe its body.
doc = Document(TEMPLATE)
body = doc.element.body
for el in list(body):
    if el.tag != qn("w:sectPr"):  # keep section (margins/page size)
        body.remove(el)


def strip_number_prefix(text):
    """Remove manual 'CHƯƠNG N.' / 'N.N.' prefixes — headings auto-number via template styles."""
    t = re.sub(r"^CHƯƠNG\s+\d+\.\s*", "", text)
    t = re.sub(r"^\d+(\.\d+)*\.\s*", "", t)
    return t.strip()


def disable_numbering(par):
    """Remove auto-numbering for a single heading paragraph (e.g. TÓM TẮT)."""
    ppr = par._p.get_or_add_pPr()
    numpr = ppr.makeelement(qn("w:numPr"), {})
    ilvl = numpr.makeelement(qn("w:ilvl"), {qn("w:val"): "0"})
    numid = numpr.makeelement(qn("w:numId"), {qn("w:val"): "0"})
    numpr.append(ilvl)
    numpr.append(numid)
    ppr.append(numpr)


def add_runs(par, text):
    """Render **bold**, *italic*, `code` markdown spans into runs."""
    for tok in re.split(r"(\*\*.+?\*\*|\*.+?\*|`.+?`)", text):
        if not tok:
            continue
        if tok.startswith("**") and tok.endswith("**"):
            par.add_run(tok[2:-2]).bold = True
        elif tok.startswith("`") and tok.endswith("`"):
            r = par.add_run(tok[1:-1])
            r.font.name = "Consolas"
        elif tok.startswith("*") and tok.endswith("*") and len(tok) > 2:
            par.add_run(tok[1:-1]).italic = True
        else:
            par.add_run(tok)


def add_body_par(text, style=None, indent=True):
    p = doc.add_paragraph(style=style)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if indent and style is None:
        p.paragraph_format.first_line_indent = Cm(1.27)  # match template body
    add_runs(p, text)
    return p


def add_caption(text):
    """Table/figure caption: template Caption style (centered, italic)."""
    p = doc.add_paragraph(style="Caption")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    add_runs(p, text)
    return p


def add_table(rows):
    cells = [[c.strip() for c in r.strip().strip("|").split("|")] for r in rows]
    cells = [r for r in cells if not all(re.fullmatch(r":?-{2,}:?", c) for c in r)]
    if not cells:
        return
    t = doc.add_table(rows=len(cells), cols=len(cells[0]))
    t.style = "Table Grid"
    for i, row in enumerate(cells):
        for j, val in enumerate(row):
            if j >= len(t.rows[i].cells):
                continue
            cell = t.rows[i].cells[j]
            cell.text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.first_line_indent = Cm(0)
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            add_runs(p, re.sub(r"\*\*(.+?)\*\*", r"\1", val) if i == 0 else val)
            for r in p.runs:
                r.font.size = Pt(12)
                if i == 0:
                    r.bold = True
    doc.add_paragraph()


def process_file(path):
    lines = path.read_text(encoding="utf-8").splitlines()
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        if not line.strip():
            i += 1
            continue
        if line.lstrip().startswith("|"):
            tbl = []
            while i < len(lines) and lines[i].lstrip().startswith("|"):
                tbl.append(lines[i])
                i += 1
            add_table(tbl)
            continue
        m = re.match(r"^(#{1,4})\s+(.*)", line)
        if m:
            level = min(len(m.group(1)), 3)
            title = strip_number_prefix(m.group(2))
            h = doc.add_heading(title, level=level)
            # abstract chapter: H1 without auto chapter number
            if level == 1 and "TÓM TẮT" in title:
                disable_numbering(h)
            i += 1
            continue
        if re.match(r"^\[Hình .+\]$", line.strip()):
            add_caption(line.strip())
            i += 1
            continue
        if re.match(r"^Bảng \d+\.\d+\s*[:.]", line.strip()):
            add_caption(line.strip())
            i += 1
            continue
        bm = re.match(r"^(\s*)[-*]\s+(.*)", line)
        if bm:
            p = doc.add_paragraph(style="List Paragraph")
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.first_line_indent = Cm(0)
            p.paragraph_format.left_indent = Cm(0.63)
            add_runs(p, "– " + bm.group(2))
            i += 1
            continue
        nm = re.match(r"^\s*(\d+)\.\s+(.*)", line)
        if nm:
            p = doc.add_paragraph(style="List Paragraph")
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.first_line_indent = Cm(0)
            p.paragraph_format.left_indent = Cm(0.63)
            add_runs(p, f"{nm.group(1)}. {nm.group(2)}")
            i += 1
            continue
        buf = [line]
        i += 1
        while i < len(lines) and lines[i].strip() and not re.match(
            r"^(#|\||[-*]\s|\s*\d+\.\s|\[Hình|Bảng \d)", lines[i].lstrip()
        ):
            buf.append(lines[i].strip())
            i += 1
        add_body_par(" ".join(buf))


# TOC note page
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.first_line_indent = Cm(0)
r = p.add_run("[MỤC LỤC — chèn Table of Contents tự động: References → Table of Contents → Update Field]")
r.italic = True
doc.add_page_break()

for idx, f in enumerate(DRAFTS):
    process_file(f)
    if idx < len(DRAFTS) - 1:
        doc.add_page_break()

doc.save(OUT)
print("Saved:", OUT)
