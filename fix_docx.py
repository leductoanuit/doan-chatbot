# -*- coding: utf-8 -*-
"""
Fix Bao_Cao_Do_An_UIT_Chatbot.docx theo feedback buoi meeting (script.md).
Cac cho them moi se duoc to do + ghi chu [GHI CHU: ...].
"""

import shutil
from docx import Document
from docx.shared import RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SRC = "Bao_Cao_Do_An_UIT_Chatbot.docx"
DST = "Bao_Cao_Do_An_UIT_Chatbot_REVISED.docx"

shutil.copy2(SRC, DST)
doc = Document(DST)


def make_red_para_xml(text, style_id=None):
    """Tao mot element <w:p> chua text mau do."""
    p = OxmlElement("w:p")
    if style_id:
        pPr = OxmlElement("w:pPr")
        pStyle = OxmlElement("w:pStyle")
        pStyle.set(qn("w:val"), style_id)
        pPr.append(pStyle)
        p.append(pPr)
    r = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "FF0000")
    rPr.append(color)
    r.append(rPr)
    t = OxmlElement("w:t")
    t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    t.text = text
    r.append(t)
    p.append(r)
    return p


def insert_after(ref_para, text, style_id=None):
    """Chen paragraph mau do ngay sau ref_para."""
    new_p = make_red_para_xml(text, style_id)
    ref_para._element.addnext(new_p)
    return new_p


def insert_before(ref_para, text, style_id=None):
    """Chen paragraph mau do ngay truoc ref_para."""
    new_p = make_red_para_xml(text, style_id)
    ref_para._element.addprevious(new_p)
    return new_p


# ─────────────────────────────────────────────────────────────────────────────
# PHAN 1: Ghi chu tong hop o dau tai lieu
# ─────────────────────────────────────────────────────────────────────────────
muc_luc_para = doc.paragraphs[0]
insert_after(muc_luc_para, (
    "=== GHI CHU TONG HOP CAC LOI FORMAT CAN SUA THU CONG (theo feedback buoi meeting) === "
    "1. Dau gach noi (en-dash): dung -- khong phai dau - ban phim thong thuong. "
    "Trong Word: Insert > Symbol > Special Characters > En Dash. "
    "2. Truoc dau '...' (ba cham) phai co dau phay: 'vi du, ...' chu khong phai 'vi du...'. "
    "3. So thap phan: dung dau phay (1,000 khong phai 1.000) theo chuan tieng Viet. "
    "4. Cac cum tu chuyen nganh KHONG in dam tru khi la heading. "
    "5. Gian dong va do cao dong trong bang: kiem tra lai toan bo bang bi sai. "
    "6. Muc luc: cap nhat lai sau khi chinh sua xong (References > Update Table). "
    "7. Dau muc bi trung (chuong 5 va chuong 6): kiem tra lai tieu de cac chuong. "
    "8. Indent bullet: cac y con phai thut vao sau hon y cha -- kiem tra muc 1.3. "
    "9. Hinh anh bi nen xau: xoa nen hoac dat wrap text = In Line with Text. "
    "10. Mui ten trong so do: them chu thich ky hieu de nguoi doc hieu duong lien/dut."
))

# ─────────────────────────────────────────────────────────────────────────────
# PHAN 2: TOM TAT -- Them y "Huong phat trien" con thieu (thay nhac 1:38-1:41)
# ─────────────────────────────────────────────────────────────────────────────
tomtat_last = doc.paragraphs[8]
insert_after(tomtat_last, (
    "[GHI CHU: BO SUNG -- Thay nhac thieu y thu 4 trong tom tat (Huong phat trien)] "
    "Ve huong phat trien, de tai de xuat cac cai tien trong tam gom: "
    "giam do tre phan hoi bang cach trien khai reranker tren GPU, "
    "cai thien do chinh xac cau tra loi qua tinh chinh prompt va few-shot examples, "
    "mo rong pham vi tri thuc sang cac don vi khac trong DHQG-HCM, "
    "va bo sung danh gia boi nguoi dung thuc te."
))

# ─────────────────────────────────────────────────────────────────────────────
# PHAN 3: MUC TIEU -- Ghi chu sap xep lai thu tu (thay nhac 3:29-4:09)
# ─────────────────────────────────────────────────────────────────────────────
muctieu_heading = doc.paragraphs[15]
insert_after(muctieu_heading, (
    "[GHI CHU: SAP XEP LAI -- Thay yeu cau sap xep muc tieu theo dung thu tu thuc hien: "
    "(1) Kien truc he thong > (2) Xay dung CSDL vector > (3) Cai dat he thong RAG "
    "> (4) Danh gia > (5) Demo. "
    "Kiem tra va chinh lai thu tu cac bullet ben duoi cho khop.]"
))

# ─────────────────────────────────────────────────────────────────────────────
# PHAN 4: PHAT BIEU BAI TOAN 1.3 -- Them vi du dau vao / dau ra (thay nhac 4:55-5:46)
# ─────────────────────────────────────────────────────────────────────────────
dau_ra_para = doc.paragraphs[24]
insert_after(dau_ra_para, (
    "[GHI CHU: BO SUNG VI DU -- Thay yeu cau bo sung vi du minh hoa dau vao/dau ra] "
    "Vi du minh hoa: Dau vao -- 'Dieu kien de duoc xet tot nghiep he tu xa UIT la gi?'; "
    "Dau ra -- 'Theo Quyet dinh 1499/QD-DHCNTT, sinh vien can hoan thanh du so tin chi "
    "theo chuong trinh dao tao, dat chuan dau ra ngoai ngu va CNTT, "
    "khong trong thoi gian bi ky luat, va nop don xet tot nghiep dung thoi han quy dinh.'"
), style_id="ListParagraph")

# ─────────────────────────────────────────────────────────────────────────────
# PHAN 5: DOI TUONG VA PHAM VI -- Them "Pham vi trien khai" (thay nhac 7:19-8:44)
# ─────────────────────────────────────────────────────────────────────────────
phamvi_last = doc.paragraphs[34]
insert_after(phamvi_last, (
    "[GHI CHU: BO SUNG -- Thay yeu cau them 'Pham vi trien khai'] "
    "-- Pham vi trien khai: he thong duoc thiet ke de trien khai cuc bo (on-premise) "
    "tai moi truong may tinh ca nhan hoac may chu noi bo cua nha truong; "
    "doi tuong su dung la sinh vien va hoc vien he dao tao tu xa UIT "
    "thong qua trinh duyet web; khong yeu cau cai dat phan mem phia nguoi dung."
), style_id="ListParagraph")

# ─────────────────────────────────────────────────────────────────────────────
# PHAN 6: Y NGHIA DE TAI -- Mo rong (thay nhac 9:17)
# ─────────────────────────────────────────────────────────────────────────────
ynghia_thuctien = doc.paragraphs[37]
insert_after(ynghia_thuctien, (
    "[GHI CHU: MO RONG -- Thay nhac phan y nghia viet qua ngan, can bo sung] "
    "Y nghia dao tao: de tai dong gop bo du lieu kiem thu (100 cau hoi -- dap an gan nhan) "
    "va bo khung danh gia tu dong theo dinh huong Auepora, co the tai su dung cho cac nghien cuu "
    "chatbot hoc vu tieng Viet ve sau. "
    "Y nghia mo rong: mo hinh trien khai va quy trinh xay dung cua de tai co the ap dung cho "
    "cac don vi dao tao khac trong he thong DHQG-HCM hoac cac truong co nhu cau tuong tu."
))

# ─────────────────────────────────────────────────────────────────────────────
# PHAN 7: CHUONG 2 -- Ghi chu bo sung trich dan (thay nhac 11:33-12:24)
# ─────────────────────────────────────────────────────────────────────────────
chuong2_heading = doc.paragraphs[47]
insert_after(chuong2_heading, (
    "[GHI CHU: BO SUNG TRICH DAN -- Thay yeu cau moi luan diem trong Chuong 2 phai co nguon "
    "theo chuan APA: (Tac gia, Nam). Neu 1 tac gia: (Ho ten, Nam); "
    "2 tac gia: (Tac gia 1 & Tac gia 2, Nam); tu 3 tac gia tro len: (Tac gia 1 va cong su, Nam). "
    "Ra soat toan bo chuong 2 va bo sung trich dan con thieu."
))

# ─────────────────────────────────────────────────────────────────────────────
# PHAN 8: CHUONG 2 -- Them bang he thong hoa NC lien quan (thay nhac 16:45-17:55)
# ─────────────────────────────────────────────────────────────────────────────
khoangtrong_heading = doc.paragraphs[62]
insert_before(khoangtrong_heading, (
    "[GHI CHU: BO SUNG BANG -- Thay yeu cau them bang he thong hoa cac nghien cuu lien quan. "
    "Mau bang gom cac cot: (1) Cong trinh [trich dan], (2) Bai toan/Ung dung, "
    "(3) Phuong phap, (4) Ket qua chinh, (5) Uu diem & Han che. "
    "Hoac dang ma tran tieu chi: cac hang la cong trinh, cac cot la tieu chi -- danh dau V neu dap ung. "
    "Bang nay nen xuat hien o cuoi muc 2.2 truoc khi vao phan Khoang trong nghien cuu."
))

# ─────────────────────────────────────────────────────────────────────────────
# PHAN 9: CHUONG 3 -- Sua "ung vien tri thuc" thanh "tai lieu tri thuc" (thay nhac 29:18-29:46)
# ─────────────────────────────────────────────────────────────────────────────
fixed_count = 0
for para in doc.paragraphs:
    for run in para.runs:
        if "ng viên" in run.text and ("tri th" in run.text or "truy xu" in run.text):
            run.text = run.text.replace(
                "Danh sách ứng viên tri thức",
                "Danh sách tài liệu tri thức [ĐÃ SỬA]"
            ).replace(
                "ứng viên sau truy xuất",
                "tài liệu ứng viên sau truy xuất [ĐÃ SỬA]"
            )
            # Set red
            rPr = run._r.get_or_add_rPr()
            color = OxmlElement("w:color")
            color.set(qn("w:val"), "FF0000")
            rPr.append(color)
            fixed_count += 1

# Ghi chu tai doan sang loc (para 115)
sang_loc_para = doc.paragraphs[115]
insert_after(sang_loc_para, (
    "[GHI CHU: SUA TU NGU -- Thay nhac 'ung vien tri thuc' de hieu nham la con nguoi. "
    "Da sua thanh 'tai lieu tri thuc' hoac 'tai lieu ung vien'. "
    f"Tim thay va sua {fixed_count} cho trong van ban."
))

# ─────────────────────────────────────────────────────────────────────────────
# PHAN 10: CHUONG 4 -- Mo ta tap kiem thu (thay nhac 36:56-37:57)
# ─────────────────────────────────────────────────────────────────────────────
kiemthu_para = doc.paragraphs[142]
insert_after(kiemthu_para, (
    "[GHI CHU: BO SUNG -- Thay yeu cau mo ta ro hon tap danh gia: "
    "(a) Tong so cau hoi va phan bo theo chu de (dao tao / tuyen sinh / hoc phi / quy che...); "
    "(b) Phan bo do kho (de / trung binh / kho); "
    "(c) Vi du minh hoa 2-3 cap cau hoi -- dap an dai dien. "
    "QUAN TRONG: Tap kiem thu phai do nhom tu xay dung, KHONG dung AI sinh toan bo "
    "vi nguoi cham se khong tin tuong do chinh xac cua tap danh gia (thay nhac ro dieu nay)."
))

# ─────────────────────────────────────────────────────────────────────────────
# PHAN 11: DANH MUC TLTKHAM KHAO -- Sua dinh dang (thay nhac 45:09-46:03)
# ─────────────────────────────────────────────────────────────────────────────
# Tim paragraph cuoi
for i in range(len(doc.paragraphs) - 1, -1, -1):
    if doc.paragraphs[i].text.strip():
        last_para = doc.paragraphs[i]
        break

insert_after(last_para, (
    "[GHI CHU: SUA DINH DANG TLTKHAM KHAO -- Thay nhac: "
    "(1) Nguon Nghi quyet 57 phai ghi la cua Nha nuoc/Chinh phu, KHONG phai trang thuvienphapluat.vn "
    "(do la nguon thu cap, khong phai nguon goc). "
    "(2) Toan bo danh muc can tuan thu dung chuan APA: "
    "Tac gia. (Nam). Ten tai lieu. Nha xuat ban/Tap chi, So(Volume), trang. "
    "(3) Ra soat lai toan bo danh muc, dac biet cac van ban phap ly can ghi dung co quan ban hanh."
))

doc.save(DST)
print(f"Da luu: {DST}")
print("Cac thay doi duoc to do va ghi chu '[GHI CHU: ...]' trong file.")
