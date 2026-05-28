"""Export short evaluation methodology summary to Word (.docx) — Times New Roman."""

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUTPUT_PATH = os.path.join(os.path.dirname(__file__), "..", "tom-tat-phuong-phap-kiem-thu.docx")
FONT = "Times New Roman"


def set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def style_run(run, size=12, bold=False, color=None, italic=False):
    run.font.name = FONT
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def add_heading(doc, text, size=13, color="1F4E79", align=WD_ALIGN_PARAGRAPH.LEFT):
    p = doc.add_paragraph()
    p.alignment = align
    run = p.add_run(text)
    style_run(run, size=size, bold=True, color=color)
    return p


def add_para(doc, text, size=12, bold=False, indent=0):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    run = p.add_run(text)
    style_run(run, size=size, bold=bold)
    return p


def add_bullet(doc, text, bold_part=None, size=12):
    p = doc.add_paragraph(style="List Bullet")
    if bold_part:
        r1 = p.add_run(bold_part)
        style_run(r1, size=size, bold=True)
        r2 = p.add_run(text)
        style_run(r2, size=size)
    else:
        r = p.add_run(text)
        style_run(r, size=size)


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    # Header
    for i, h in enumerate(headers):
        cell = table.cell(0, i)
        set_cell_bg(cell, "1F4E79")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        style_run(run, size=11, bold=True, color="FFFFFF")
    # Rows
    for ri, row in enumerate(rows):
        bg = "DCE6F1" if ri % 2 == 0 else "FFFFFF"
        for ci, val in enumerate(row):
            cell = table.cell(ri + 1, ci)
            set_cell_bg(cell, bg)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if ci > 0 else WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(str(val))
            style_run(run, size=11)
    doc.add_paragraph()


def build():
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3)
        section.right_margin = Cm(2)

    # Title
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("PHƯƠNG PHÁP KIỂM THỬ HỆ THỐNG RAG CHATBOT")
    style_run(r, size=14, bold=True, color="1F4E79")

    t2 = doc.add_paragraph()
    t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = t2.add_run("Tư Vấn Đào Tạo UIT — Theo Auepora Framework")
    style_run(r2, size=12, italic=True, color="595959")

    doc.add_paragraph()

    # 1. Tổng quan
    add_heading(doc, "1. Tổng Quan")
    add_para(doc,
        "Hệ thống được kiểm thử theo Auepora Framework (paper 2405.07437) — "
        "bộ tiêu chí đánh giá RAG toàn diện gồm 3 câu hỏi cốt lõi: What to Evaluate? "
        "How to Evaluate? How to Quantify? Kiểm thử chia làm 2 tầng độc lập.")

    doc.add_paragraph()

    # 2. Tầng 1
    add_heading(doc, "2. Tầng 1 — Unit Tests (Retrieval Metrics)")
    add_para(doc, "Kiểm tra tính đúng đắn về mặt toán học của 6 hàm tính retrieval metrics, "
             "không phụ thuộc LLM hay mạng, chạy dưới 0.1 giây.")
    doc.add_paragraph()
    add_table(doc,
        ["Metric", "Ý nghĩa"],
        [
            ["Precision@K", "Tỉ lệ chunk liên quan trong top-K kết quả trả về"],
            ["Recall@K", "Tỉ lệ chunk đúng được tìm thấy trên tổng chunk liên quan"],
            ["Hit@K", "Có ít nhất 1 chunk đúng trong top-K không"],
            ["MRR", "Chunk đúng đầu tiên xuất hiện ở vị trí nào trong ranking"],
            ["AP@K", "Average Precision cho 1 query tại cutoff K"],
            ["MAP@K", "Trung bình AP@K trên toàn bộ tập query"],
        ]
    )
    add_para(doc, "Mỗi metric kiểm tra 6 trường hợp: happy path, all-hits (=1.0), "
             "no-hits (=0.0), empty input, k > độ dài list, và ví dụ từ sách giáo khoa IR.")
    add_para(doc, "Kết quả: 36/36 test cases PASS.", bold=True)

    doc.add_paragraph()

    # 3. Tầng 2
    add_heading(doc, "3. Tầng 2 — RAG Evaluation (LLM-as-Judge)")
    add_para(doc,
        "Đánh giá chất lượng đầu cuối bằng cách chạy từng câu hỏi qua RAG pipeline thật, "
        "sau đó dùng Gemini API làm judge chấm điểm 0.0–1.0 cho từng metric theo prompt tiếng Việt.")

    doc.add_paragraph()
    add_heading(doc, "3.1 Dataset", size=12, color="2E75B6")
    add_table(doc,
        ["File", "Số câu", "Loại"],
        [
            ["eval-dataset.json", "35", "Câu hỏi thật về UIT (tuyển sinh, học vụ, học phí, quy chế, CTĐT)"],
            ["eval-dataset-robustness.json", "11", "5 câu noise + 6 câu out-of-domain"],
            ["Tổng", "46", ""],
        ]
    )

    add_heading(doc, "3.2 Metrics đánh giá", size=12, color="2E75B6")
    add_table(doc,
        ["Nhóm", "Metric", "Câu hỏi đánh giá"],
        [
            ["Generation", "Faithfulness", "Answer có bịa thông tin ngoài context không?"],
            ["Generation", "Answer Relevancy", "Answer có trả lời đúng câu hỏi không?"],
            ["Generation", "Correctness", "Answer có khớp với ground truth không?"],
            ["Retrieval", "Context Precision", "Chunk retrieve có liên quan không?"],
            ["Retrieval", "Context Recall", "Chunk cần thiết có được tìm thấy không?"],
            ["Robustness", "Noise Robustness", "Answer vẫn đúng khi context bị nhiễu?"],
            ["Robustness", "Negative Rejection", "Hệ thống từ chối câu hỏi ngoài domain?"],
        ]
    )

    add_heading(doc, "3.3 Phân loại theo eval_type", size=12, color="2E75B6")
    add_table(doc,
        ["eval_type", "Số câu", "Metrics chấm"],
        [
            ["standard", "35", "5 metrics: faithfulness, relevancy, precision, recall, correctness"],
            ["noise", "5", "5 metrics trên + noise_robustness"],
            ["negative", "6", "Chỉ negative_rejection"],
        ]
    )

    doc.add_paragraph()

    # 4. Cơ sở lý thuyết
    add_heading(doc, "4. Cơ Sở Lý Thuyết — LLM-as-Judge")
    add_para(doc,
        "Theo Auepora paper (Table 1), LLM-as-Judge là phương pháp được áp dụng phổ biến nhất "
        "trong các framework đánh giá RAG hiện đại (TruEra, RAGAs, LangChain, DomainRAG...). "
        "Paper khẳng định: "
        "\"The approach of employing LLMs as evaluative judges is a versatile and automatic "
        "method for quality assessment, catering to instances where traditional ground truths "
        "may be elusive.\"")
    doc.add_paragraph()
    add_para(doc, "Biện pháp giảm thiểu bias trong hệ thống:", bold=True)
    add_bullet(doc, "Dùng Gemini judge nhưng chatbot dùng DeepSeek — tránh self-evaluation bias")
    add_bullet(doc, "Có ground_truth do con người viết để neo điểm correctness")
    add_bullet(doc, "negative_rejection dùng regex pattern — không phụ thuộc LLM")

    doc.add_paragraph()

    # 5. Tóm tắt
    add_heading(doc, "5. Tóm Tắt So Sánh")
    add_table(doc,
        ["", "Tầng 1 — Unit Tests", "Tầng 2 — RAG Evaluation"],
        [
            ["Công cụ", "Python unittest", "Gemini LLM-as-Judge"],
            ["Thời gian", "0.075 giây", "~3–4 phút"],
            ["Số test", "36 test cases", "46 câu hỏi"],
            ["Mục đích", "Đúng toán học", "Đúng ngữ nghĩa"],
            ["Phụ thuộc", "Không cần gì", "Qdrant + Gemini API"],
        ]
    )

    output = os.path.abspath(OUTPUT_PATH)
    doc.save(output)
    print(f"Saved: {output}")


if __name__ == "__main__":
    build()
