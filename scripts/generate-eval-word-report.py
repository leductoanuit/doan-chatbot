"""Generate Word report from eval summary JSON files."""

import json
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


RESULTS_DIR = Path("tests/evaluation/results")
OUTPUT_PATH = RESULTS_DIR / "bao-cao-ket-qua-eval.docx"

BATCH1_SUMMARY = RESULTS_DIR / "eval-summary-20260602_173030.json"
BATCH2_SUMMARY = RESULTS_DIR / "eval-summary-20260602_193653.json"

METRIC_LABELS = {
    "faithfulness": "Độ trung thực (Faithfulness)",
    "answer_relevancy": "Độ liên quan câu trả lời (Answer Relevancy)",
    "context_precision": "Độ chính xác ngữ cảnh (Context Precision)",
    "context_recall": "Độ bao phủ ngữ cảnh (Context Recall)",
    "correctness": "Độ đúng đắn (Correctness)",
}

CATEGORY_LABELS = {
    "ctdt": "Chương trình đào tạo (CTĐT)",
    "hoc_phi": "Học phí",
    "hoc_vu": "Học vụ",
    "quy_che": "Quy chế",
    "tuyen_sinh": "Tuyển sinh",
}

DIFFICULTY_LABELS = {
    "easy": "Dễ",
    "medium": "Trung bình",
    "hard": "Khó",
}


def set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def set_cell_border(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        border = OxmlElement(f"w:{side}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:color"), "CCCCCC")
        tcBorders.append(border)
    tcPr.append(tcBorders)


def score_color(score: float) -> str:
    if score >= 0.85:
        return "C6EFCE"  # green
    elif score >= 0.70:
        return "FFEB9C"  # yellow
    else:
        return "FFC7CE"  # red


def add_heading(doc: Document, text: str, level: int):
    h = doc.add_heading(text, level=level)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return h


def add_metrics_table(doc: Document, metrics: dict, title: str):
    doc.add_paragraph(title).runs[0].bold = True

    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header
    hdr = table.rows[0].cells
    for cell, text in zip(hdr, ["Metric", "Điểm", "Đánh giá"]):
        cell.text = text
        cell.paragraphs[0].runs[0].bold = True
        set_cell_bg(cell, "4472C4")
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    for key, label in METRIC_LABELS.items():
        if key not in metrics:
            continue
        score = metrics[key]
        row = table.add_row().cells
        row[0].text = label
        row[1].text = f"{score:.4f}"
        row[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        row[2].text = "Tốt" if score >= 0.85 else ("Khá" if score >= 0.70 else "Cần cải thiện")
        row[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        for cell in row:
            set_cell_bg(cell, score_color(score))
            set_cell_border(cell)

    doc.add_paragraph()


def add_category_table(doc: Document, by_category: dict, title: str):
    doc.add_paragraph(title).runs[0].bold = True

    metrics_keys = ["faithfulness", "answer_relevancy", "context_precision", "context_recall"]
    short_labels = ["Faithfulness", "Ans. Relevancy", "Ctx. Precision", "Ctx. Recall"]

    table = doc.add_table(rows=1, cols=len(metrics_keys) + 1)
    table.style = "Table Grid"

    hdr = table.rows[0].cells
    hdr[0].text = "Danh mục"
    hdr[0].paragraphs[0].runs[0].bold = True
    set_cell_bg(hdr[0], "4472C4")
    hdr[0].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    for i, lbl in enumerate(short_labels):
        hdr[i + 1].text = lbl
        hdr[i + 1].paragraphs[0].runs[0].bold = True
        hdr[i + 1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_bg(hdr[i + 1], "4472C4")
        hdr[i + 1].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    for cat_key, cat_data in by_category.items():
        row = table.add_row().cells
        row[0].text = CATEGORY_LABELS.get(cat_key, cat_key)
        set_cell_border(row[0])
        for i, mk in enumerate(metrics_keys):
            score = cat_data.get(mk, 0)
            row[i + 1].text = f"{score:.4f}"
            row[i + 1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_cell_bg(row[i + 1], score_color(score))
            set_cell_border(row[i + 1])

    doc.add_paragraph()


def add_difficulty_table(doc: Document, by_difficulty: dict, title: str):
    doc.add_paragraph(title).runs[0].bold = True

    metrics_keys = ["faithfulness", "answer_relevancy", "context_precision", "context_recall"]
    short_labels = ["Faithfulness", "Ans. Relevancy", "Ctx. Precision", "Ctx. Recall"]

    table = doc.add_table(rows=1, cols=len(metrics_keys) + 1)
    table.style = "Table Grid"

    hdr = table.rows[0].cells
    hdr[0].text = "Độ khó"
    hdr[0].paragraphs[0].runs[0].bold = True
    set_cell_bg(hdr[0], "4472C4")
    hdr[0].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    for i, lbl in enumerate(short_labels):
        hdr[i + 1].text = lbl
        hdr[i + 1].paragraphs[0].runs[0].bold = True
        hdr[i + 1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_bg(hdr[i + 1], "4472C4")
        hdr[i + 1].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    for diff_key in ["easy", "medium", "hard"]:
        if diff_key not in by_difficulty:
            continue
        diff_data = by_difficulty[diff_key]
        row = table.add_row().cells
        row[0].text = DIFFICULTY_LABELS.get(diff_key, diff_key)
        set_cell_border(row[0])
        for i, mk in enumerate(metrics_keys):
            score = diff_data.get(mk, 0)
            row[i + 1].text = f"{score:.4f}"
            row[i + 1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_cell_bg(row[i + 1], score_color(score))
            set_cell_border(row[i + 1])

    doc.add_paragraph()


def add_latency_table(doc: Document, latency: dict, title: str):
    doc.add_paragraph(title).runs[0].bold = True

    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"

    hdr = table.rows[0].cells
    for cell, text in zip(hdr, ["Chỉ số", "Giá trị"]):
        cell.text = text
        cell.paragraphs[0].runs[0].bold = True
        set_cell_bg(cell, "4472C4")
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    latency_labels = {
        "avg_ms": "Trung bình (avg)",
        "p50_ms": "Trung vị (P50)",
        "p90_ms": "P90",
        "p99_ms": "P99",
        "min_ms": "Nhỏ nhất",
        "max_ms": "Lớn nhất",
    }

    for key, label in latency_labels.items():
        if key not in latency:
            continue
        val_ms = latency[key]
        row = table.add_row().cells
        row[0].text = label
        row[1].text = f"{val_ms / 1000:.1f}s"
        row[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        for cell in row:
            set_cell_border(cell)

    doc.add_paragraph()


def add_rank_metrics_table(doc: Document, rr: dict, title: str):
    if not rr:
        return
    doc.add_paragraph(title).runs[0].bold = True

    k = rr.get("k", 10)
    rows_data = [
        (f"MAP@{k}", rr.get(f"map_at_{k}", 0)),
        ("MRR", rr.get("mrr", 0)),
        (f"Hit@{k}", rr.get(f"hit_at_{k}", 0)),
    ]

    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for cell, text in zip(hdr, ["Metric", "Điểm", "Đánh giá"]):
        cell.text = text
        cell.paragraphs[0].runs[0].bold = True
        set_cell_bg(cell, "375623")
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    for label, score in rows_data:
        row = table.add_row().cells
        row[0].text = label
        row[1].text = f"{score:.4f}"
        row[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        row[2].text = "Tốt" if score >= 0.75 else ("Khá" if score >= 0.55 else "Cần cải thiện")
        row[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        for cell in row:
            set_cell_bg(cell, score_color(score))
            set_cell_border(cell)

    # Per-category rank metrics
    by_cat = rr.get("by_category", {})
    if by_cat:
        doc.add_paragraph()
        p = doc.add_paragraph(f"Retrieval Rank Metrics theo danh mục (k={k}):")
        p.runs[0].bold = True
        cat_table = doc.add_table(rows=1, cols=4)
        cat_table.style = "Table Grid"
        hdr2 = cat_table.rows[0].cells
        for cell, text in zip(hdr2, ["Danh mục", f"MAP@{k}", "MRR", f"Hit@{k}"]):
            cell.text = text
            cell.paragraphs[0].runs[0].bold = True
            set_cell_bg(cell, "375623")
            cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        for cat_key, cat_data in by_cat.items():
            row = cat_table.add_row().cells
            row[0].text = CATEGORY_LABELS.get(cat_key, cat_key)
            set_cell_border(row[0])
            for i, mk in [("map_at_k", 1), ("mrr", 2), ("hit_at_k", 3)]:
                score = cat_data.get(i, 0)
                row[mk].text = f"{score:.4f}"
                row[mk].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                set_cell_bg(row[mk], score_color(score))
                set_cell_border(row[mk])

    doc.add_paragraph()


def compute_combined(s1: dict, s2: dict) -> dict:
    """Average metrics from two equal-sized batches."""
    combined = {}
    for key in s1["metrics"]:
        combined[key] = round((s1["metrics"][key] + s2["metrics"].get(key, 0)) / 2, 4)
    return combined


def main():
    with open(BATCH1_SUMMARY, encoding="utf-8") as f:
        s1 = json.load(f)
    with open(BATCH2_SUMMARY, encoding="utf-8") as f:
        s2 = json.load(f)

    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3)
        section.right_margin = Cm(2)

    # Title
    title = doc.add_heading("BÁO CÁO KẾT QUẢ ĐÁNH GIÁ HỆ THỐNG RAG CHATBOT", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph("Trường Đại học Công nghệ Thông tin – CITD")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].bold = True

    date_p = doc.add_paragraph(f"Ngày thực hiện: 02/06/2026  |  Mô hình đánh giá: gemini-2.5-flash")
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # 1. Tổng quan
    add_heading(doc, "1. TỔNG QUAN", 1)
    overview_data = [
        ("Tổng số câu hỏi đánh giá", "100 câu (2 batch × 50 câu)"),
        ("Tập dữ liệu", "eval-dataset.json (100 câu)"),
        ("Mô hình embedding", "BAAI/bge-m3"),
        ("Mô hình reranker", "BAAI/bge-reranker-v2-m3"),
        ("Mô hình LLM (RAG)", "gemini-2.5-flash"),
        ("Mô hình LLM judge", "gemini-2.5-flash"),
        ("Vector store", "Qdrant (513 chunks)"),
        ("Metadata store", "PostgreSQL (19 nguồn tài liệu)"),
        ("Phương pháp đánh giá", "RAGAS framework (faithfulness, answer relevancy, context precision/recall, correctness)"),
    ]

    table = doc.add_table(rows=len(overview_data), cols=2)
    table.style = "Table Grid"
    for i, (key, val) in enumerate(overview_data):
        row = table.rows[i].cells
        row[0].text = key
        row[0].paragraphs[0].runs[0].bold = True
        row[1].text = val
        set_cell_bg(row[0], "DDEEFF")
        for cell in row:
            set_cell_border(cell)

    doc.add_paragraph()

    # 2. Kết quả tổng hợp
    add_heading(doc, "2. KẾT QUẢ TỔNG HỢP (100 CÂU)", 1)

    combined = compute_combined(s1, s2)
    add_metrics_table(doc, combined, "2.1 Các metric tổng hợp (trung bình 2 batch):")

    # 3. Kết quả theo batch
    add_heading(doc, "3. KẾT QUẢ THEO BATCH", 1)

    add_heading(doc, "3.1 Batch 1 – Câu 1 đến 50", 2)
    add_metrics_table(doc, s1["metrics"], "Metric tổng thể:")
    add_rank_metrics_table(doc, s1.get("retrieval_rank_metrics", {}), "Retrieval Rank Metrics (Auepora):")
    add_category_table(doc, s1["by_category"], "Theo danh mục:")
    add_difficulty_table(doc, s1["by_difficulty"], "Theo độ khó:")
    add_latency_table(doc, s1["latency"], "Thời gian phản hồi:")

    add_heading(doc, "3.2 Batch 2 – Câu 51 đến 100", 2)
    add_metrics_table(doc, s2["metrics"], "Metric tổng thể:")
    add_rank_metrics_table(doc, s2.get("retrieval_rank_metrics", {}), "Retrieval Rank Metrics (Auepora):")
    add_category_table(doc, s2["by_category"], "Theo danh mục:")
    add_difficulty_table(doc, s2["by_difficulty"], "Theo độ khó:")
    add_latency_table(doc, s2["latency"], "Thời gian phản hồi:")

    # 4. So sánh 2 batch
    add_heading(doc, "4. SO SÁNH 2 BATCH", 1)

    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for cell, text in zip(hdr, ["Metric", "Batch 1 (Q1-50)", "Batch 2 (Q51-100)", "Chênh lệch"]):
        cell.text = text
        cell.paragraphs[0].runs[0].bold = True
        set_cell_bg(cell, "4472C4")
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    for key, label in METRIC_LABELS.items():
        v1 = s1["metrics"].get(key, 0)
        v2 = s2["metrics"].get(key, 0)
        diff = v2 - v1
        row = table.add_row().cells
        row[0].text = label
        row[1].text = f"{v1:.4f}"
        row[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        row[2].text = f"{v2:.4f}"
        row[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        row[3].text = f"{'+' if diff >= 0 else ''}{diff:.4f}"
        row[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_bg(row[3], "C6EFCE" if diff >= 0 else "FFC7CE")
        for cell in row:
            set_cell_border(cell)

    doc.add_paragraph()

    # Rank metrics comparison
    rr1 = s1.get("retrieval_rank_metrics", {})
    rr2 = s2.get("retrieval_rank_metrics", {})
    if rr1 and rr2:
        k = rr1.get("k", 10)
        doc.add_paragraph("4.2 So sánh Retrieval Rank Metrics:").runs[0].bold = True
        rank_table = doc.add_table(rows=1, cols=4)
        rank_table.style = "Table Grid"
        hdr2 = rank_table.rows[0].cells
        for cell, text in zip(hdr2, ["Metric", "Batch 1", "Batch 2", "Chênh lệch"]):
            cell.text = text
            cell.paragraphs[0].runs[0].bold = True
            set_cell_bg(cell, "375623")
            cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        rank_rows = [
            (f"MAP@{k}", f"map_at_{k}"),
            ("MRR", "mrr"),
            (f"Hit@{k}", f"hit_at_{k}"),
        ]
        for label, key in rank_rows:
            v1 = rr1.get(key, 0)
            v2 = rr2.get(key, 0)
            diff = v2 - v1
            row = rank_table.add_row().cells
            row[0].text = label
            row[1].text = f"{v1:.4f}"
            row[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            row[2].text = f"{v2:.4f}"
            row[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            row[3].text = f"{'+' if diff >= 0 else ''}{diff:.4f}"
            row[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_cell_bg(row[3], "C6EFCE" if diff >= 0 else "FFC7CE")
            for cell in row:
                set_cell_border(cell)

    doc.add_paragraph()

    # 5. Nhận xét
    add_heading(doc, "5. NHẬN XÉT VÀ ĐỀ XUẤT", 1)

    doc.add_paragraph("5.1 Điểm mạnh:").runs[0].bold = True
    strengths = [
        "Context Precision batch 1 đạt 0.94 – hệ thống retrieval hoạt động rất tốt với 50 câu đầu.",
        "Context Recall ổn định ở mức cao (0.87–0.92) cho thấy hệ thống tìm được đủ ngữ cảnh liên quan.",
        "Faithfulness và Answer Relevancy đồng đều (~0.73) giữa 2 batch.",
    ]
    for s in strengths:
        p = doc.add_paragraph(s, style="List Bullet")

    doc.add_paragraph("5.2 Điểm cần cải thiện:").runs[0].bold = True
    weaknesses = [
        "Correctness thấp (0.58–0.59) ở cả 2 batch – LLM sinh câu trả lời chưa đủ chính xác dù retrieval tốt.",
        "Context Precision batch 2 giảm xuống 0.75 (so với 0.94 batch 1) – 50 câu sau khó hơn về retrieval.",
        "Danh mục hoc_phi batch 1 có điểm thấp (0.4 faithfulness, 0.2 answer_relevancy) – cần bổ sung dữ liệu học phí.",
        "Latency cao (avg ~6.7 phút/câu batch 1, ~2.2 phút/câu batch 2) do reranker chạy trên CPU.",
    ]
    for w in weaknesses:
        p = doc.add_paragraph(w, style="List Bullet")

    doc.add_paragraph("5.3 Đề xuất cải thiện:").runs[0].bold = True
    suggestions = [
        "Cải thiện prompt LLM để tăng correctness – có thể dùng chain-of-thought hoặc few-shot examples.",
        "Bổ sung dữ liệu học phí chi tiết hơn để cải thiện danh mục hoc_phi.",
        "Triển khai reranker trên GPU để giảm latency xuống dưới 10 giây/câu.",
        "Xem xét fine-tune hoặc RAG fusion để cải thiện context precision cho câu hỏi khó.",
    ]
    for s in suggestions:
        p = doc.add_paragraph(s, style="List Bullet")

    doc.add_paragraph()

    # Footer note
    footer_p = doc.add_paragraph(f"Báo cáo được tạo tự động ngày 04/06/2026")
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_p.runs[0].font.size = Pt(9)
    footer_p.runs[0].font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    doc.save(OUTPUT_PATH)
    print(f"Saved: {OUTPUT_PATH}")


if __name__ == "__main__":
    main()
