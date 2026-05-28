"""Export evaluation methodology slides to Word document (.docx)."""

from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUTPUT_PATH = os.path.join(os.path.dirname(__file__), "..", "phuong-phap-kiem-thu-rag-chatbot-uit.docx")


def set_cell_bg(cell, hex_color: str):
    """Set background color of a table cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def add_heading(doc: Document, text: str, level: int = 1, color: str = "1F4E79"):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in p.runs:
        run.font.color.rgb = RGBColor.from_string(color)
    return p


def add_slide_divider(doc: Document, slide_num: int, title: str):
    """Add a colored slide header bar."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_bg(cell, "1F4E79")
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(f"  Slide {slide_num}: {title}")
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    run.font.bold = True
    run.font.size = Pt(14)
    doc.add_paragraph()


def add_metric_table(doc: Document, headers: list, rows: list, col_widths: list = None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    # Header row
    for i, h in enumerate(headers):
        cell = table.cell(0, i)
        set_cell_bg(cell, "2E75B6")
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(10)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Data rows
    for ri, row in enumerate(rows):
        bg = "DEEAF1" if ri % 2 == 0 else "FFFFFF"
        for ci, val in enumerate(row):
            cell = table.cell(ri + 1, ci)
            set_cell_bg(cell, bg)
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(10)
    doc.add_paragraph()


def add_bullet(doc: Document, text: str, bold_prefix: str = None, level: int = 0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Cm(level * 0.5)
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.font.bold = True
        run.font.size = Pt(11)
        p.add_run(text).font.size = Pt(11)
    else:
        run = p.add_run(text)
        run.font.size = Pt(11)


def add_code_block(doc: Document, code: str):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    run = p.add_run(code)
    run.font.name = "Courier New"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)


def build_doc():
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # Title page
    title = doc.add_heading("Phương Pháp Kiểm Thử", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.color.rgb = RGBColor.from_string("1F4E79")
        run.font.size = Pt(24)

    sub = doc.add_paragraph("RAG Chatbot Tư Vấn Đào Tạo UIT")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].font.size = Pt(16)
    sub.runs[0].font.bold = True
    sub.runs[0].font.color.rgb = RGBColor.from_string("2E75B6")

    framework = doc.add_paragraph("Theo Auepora Framework (paper 2405.07437)")
    framework.alignment = WD_ALIGN_PARAGRAPH.CENTER
    framework.runs[0].font.size = Pt(12)
    framework.runs[0].font.color.rgb = RGBColor.from_string("7F7F7F")

    doc.add_page_break()

    # --- Slide 1: Giới thiệu ---
    add_slide_divider(doc, 1, "Giới Thiệu")
    add_heading(doc, "Mục tiêu & Phạm vi", 2)
    add_bullet(doc, "Đánh giá toàn diện chất lượng hệ thống RAG Chatbot tư vấn đào tạo UIT", bold_prefix="Mục tiêu: ")
    add_bullet(doc, "Auepora Framework (paper 2405.07437) — bộ tiêu chí đánh giá RAG có hệ thống", bold_prefix="Dựa trên: ")
    doc.add_paragraph()
    add_heading(doc, "2 Tầng Kiểm Thử", 2)
    add_bullet(doc, "Unit Tests — kiểm tra thuần toán học, không cần LLM, chạy < 1 giây", bold_prefix="Tầng 1: ")
    add_bullet(doc, "RAG Evaluation — LLM-as-Judge với 46 câu hỏi thực tế, ~3-4 phút", bold_prefix="Tầng 2: ")

    doc.add_page_break()

    # --- Slide 2: Dataset ---
    add_slide_divider(doc, 2, "Dataset — Dữ Liệu Đánh Giá")
    add_metric_table(doc,
        ["File", "Số câu", "Mục đích"],
        [
            ["eval-dataset.json", "35 câu", "Câu hỏi thật về UIT (tuyển sinh, học vụ, học phí...)"],
            ["eval-dataset-robustness.json", "11 câu", "Kiểm tra độ bền vững (noise + out-of-domain)"],
            ["Tổng", "46 câu", ""],
        ]
    )
    add_heading(doc, "Phân loại 35 câu chính", 2)
    add_metric_table(doc,
        ["Category", "Nội dung"],
        [
            ["tuyen_sinh", "Thông tin xét tuyển, hồ sơ, điều kiện"],
            ["hoc_vu", "Thủ tục học vụ, bảo lưu, thôi học"],
            ["hoc_phi", "Học phí, lệ phí, hoàn trả"],
            ["quy_che", "Quy chế, kỷ luật thi cử"],
            ["ctdt", "Chương trình đào tạo, tín chỉ"],
        ]
    )
    add_heading(doc, "11 câu Robustness", 2)
    add_bullet(doc, "5 câu noise — context bị nhiễu thông tin sai")
    add_bullet(doc, "6 câu out-of-domain — nấu phở, thời tiết, giá vàng, ChatGPT...")

    doc.add_page_break()

    # --- Slide 3: Unit Tests ---
    add_slide_divider(doc, 3, "Tầng 1 — Unit Tests (Retrieval Metrics)")
    add_heading(doc, "Đặc điểm", 2)
    add_bullet(doc, "Không cần LLM, không cần network, không cần Qdrant")
    add_bullet(doc, "Chạy trong 0.075 giây — 36 test cases, pass 100%")
    add_bullet(doc, "Dùng Python unittest, reference values tính tay từ sách giáo khoa IR")
    doc.add_paragraph()
    add_heading(doc, "6 Metrics được kiểm tra", 2)
    add_metric_table(doc,
        ["Metric", "Công thức", "Ý nghĩa"],
        [
            ["Precision@K", "|relevant ∩ top-K| / K", "Tỉ lệ chunk liên quan trong top-K kết quả"],
            ["Recall@K", "|relevant ∩ top-K| / |relevant|", "Tìm được bao nhiêu % chunk đúng"],
            ["Hit@K", "1 nếu có ≥1 chunk đúng trong K", "Có tìm được gì không?"],
            ["MRR", "1 / rank đầu tiên đúng", "Chunk đúng xuất hiện sớm không?"],
            ["AP@K", "Trung bình Precision tại vị trí đúng", "Chất lượng ranking cho 1 query"],
            ["MAP@K", "Trung bình AP@K trên nhiều query", "Chất lượng ranking tổng thể"],
        ]
    )
    add_heading(doc, "Mỗi metric kiểm tra các trường hợp", 2)
    add_bullet(doc, "Happy path — giá trị trung bình thực tế")
    add_bullet(doc, "All-hits — expected = 1.0")
    add_bullet(doc, "No-hits — expected = 0.0")
    add_bullet(doc, "Empty inputs — expected = 0.0")
    add_bullet(doc, "k > len(retrieved) — graceful behavior")

    doc.add_page_break()

    # --- Slide 4: RAG Evaluation Pipeline ---
    add_slide_divider(doc, 4, "Tầng 2 — RAG Evaluation Pipeline")
    add_heading(doc, "Luồng xử lý", 2)
    add_code_block(doc,
        "Dataset (46 câu)\n"
        "     ↓\n"
        "RAG Pipeline chạy thật (Qdrant + BGE Reranker)\n"
        "     ↓ answer + contexts\n"
        "LLM-as-Judge (Gemini API)\n"
        "     ↓ điểm 0.0 → 1.0 mỗi metric\n"
        "Báo cáo tổng hợp theo category"
    )
    doc.add_paragraph()
    add_heading(doc, "Thành phần tham gia", 2)
    add_metric_table(doc,
        ["Thành phần", "Vai trò"],
        [
            ["BAAI/bge-m3", "Embedding model — chuyển câu hỏi thành vector"],
            ["Qdrant", "Vector database — tìm kiếm chunk liên quan"],
            ["BAAI/bge-reranker-v2-m3", "Reranker — sắp xếp lại kết quả theo độ liên quan"],
            ["DeepSeek LLM", "Sinh câu trả lời từ context"],
            ["Gemini API", "Judge — chấm điểm chất lượng câu trả lời"],
        ]
    )

    doc.add_page_break()

    # --- Slide 5: 7 Metrics ---
    add_slide_divider(doc, 5, "7 Metrics Đánh Giá")
    add_metric_table(doc,
        ["Nhóm", "Metric", "Câu hỏi đánh giá"],
        [
            ["Generation", "Faithfulness", "Answer có bịa thông tin không có trong context không?"],
            ["Generation", "Answer Relevancy", "Answer có trả lời đúng câu hỏi không?"],
            ["Generation", "Correctness", "Answer có khớp với ground truth không?"],
            ["Retrieval", "Context Precision", "Chunk được retrieve có liên quan không?"],
            ["Retrieval", "Context Recall", "Chunk cần thiết có được retrieve không?"],
            ["Robustness", "Noise Robustness", "Answer vẫn đúng dù context bị nhiễu?"],
            ["Robustness", "Negative Rejection", "Hệ thống có từ chối câu hỏi ngoài domain không?"],
        ]
    )
    doc.add_paragraph()
    add_heading(doc, "Phân loại theo eval_type", 2)
    add_metric_table(doc,
        ["eval_type", "Metrics chấm"],
        [
            ["standard (35 câu)", "faithfulness + answer_relevancy + context_precision + context_recall + correctness"],
            ["noise (5 câu)", "5 metrics trên + noise_robustness"],
            ["negative (6 câu)", "Chỉ negative_rejection"],
        ]
    )

    doc.add_page_break()

    # --- Slide 6: Cách chấm ---
    add_slide_divider(doc, 6, "Cách Chấm Điểm — LLM-as-Judge")
    add_heading(doc, "Cơ chế hoạt động", 2)
    add_bullet(doc, "Gemini nhận prompt tiếng Việt mô tả tiêu chí chấm")
    add_bullet(doc, "Gemini trả về một số thập phân 0.0 → 1.0")
    add_bullet(doc, "Lấy trung bình toàn bộ dataset cho từng metric")
    doc.add_paragraph()
    add_heading(doc, "Ví dụ prompt chấm Faithfulness", 2)
    add_code_block(doc,
        "CONTEXTS: [các đoạn văn được retrieve từ Qdrant]\n"
        "CÂU HỎI: Học phí hệ đào tạo từ xa là bao nhiêu?\n"
        "CÂU TRẢ LỜI: [answer của chatbot]\n\n"
        "Đánh giá FAITHFULNESS:\n"
        "  - 1.0: mọi luận điểm đều có trong contexts\n"
        "  - 0.5: khoảng một nửa luận điểm có trong contexts\n"
        "  - 0.0: câu trả lời chứa thông tin không có trong contexts\n\n"
        "Chỉ trả lời một số thập phân từ 0.0 đến 1.0."
    )
    doc.add_paragraph()
    add_heading(doc, "Lưu ý về LLM-as-Judge", 2)
    add_bullet(doc, "LLM bias — có xu hướng chấm cao cho văn phong giống mình tạo ra")
    add_bullet(doc, "Không nhất quán — cùng 1 câu chạy 2 lần có thể ra điểm khác")
    add_bullet(doc, "Giảm thiểu: dùng Gemini judge nhưng chatbot dùng DeepSeek (khác model)")
    add_bullet(doc, "negative_rejection dùng regex pattern — không phụ thuộc LLM")

    doc.add_page_break()

    # --- Slide 7: Output ---
    add_slide_divider(doc, 7, "Kết Quả & Báo Cáo")
    add_heading(doc, "2 file output sau mỗi lần chạy", 2)
    add_bullet(doc, "eval-report-{timestamp}.json — chi tiết điểm từng câu hỏi", bold_prefix="")
    add_bullet(doc, "eval-summary-{timestamp}.json — tổng hợp aggregate + by category", bold_prefix="")
    doc.add_paragraph()
    add_heading(doc, "Mẫu kết quả tổng hợp", 2)
    add_code_block(doc,
        "Metrics:\n"
        "  faithfulness        : 0.82\n"
        "  answer_relevancy    : 0.79\n"
        "  context_precision   : 0.71\n"
        "  context_recall      : 0.68\n"
        "  correctness         : 0.75\n"
        "  noise_robustness    : 0.70\n"
        "  negative_rejection  : 0.90\n\n"
        "By category:\n"
        "  tuyen_sinh   : avg=0.80\n"
        "  hoc_vu       : avg=0.74\n"
        "  hoc_phi      : avg=0.76\n"
        "  out_of_domain: avg=0.88"
    )
    doc.add_paragraph()
    add_heading(doc, "So sánh trước/sau Auepora", 2)
    add_metric_table(doc,
        ["Phiên bản", "Ngày", "Số câu", "Số metrics"],
        [
            ["Trước Auepora", "25/04/2026", "35", "4 metrics"],
            ["Sau Auepora", "09/05/2026", "46", "7 metrics"],
        ]
    )

    doc.add_page_break()

    # --- Slide 8: Tóm tắt ---
    add_slide_divider(doc, 8, "Tóm Tắt")
    add_metric_table(doc,
        ["", "Tầng 1 — Unit Tests", "Tầng 2 — RAG Evaluation"],
        [
            ["Công cụ", "Python unittest", "Gemini LLM-as-Judge"],
            ["Thời gian", "0.075 giây", "~3-4 phút"],
            ["Số test", "36 test cases", "46 câu hỏi"],
            ["Mục đích", "Đúng toán học", "Đúng ngữ nghĩa"],
            ["Phụ thuộc", "Không cần gì", "Qdrant + Gemini API"],
        ]
    )
    doc.add_paragraph()
    add_heading(doc, "Triết lý thiết kế", 2)
    add_bullet(doc, "Tách bạch kiểm tra logic thuần túy (unit) với kiểm tra chất lượng thực tế (eval)")
    add_bullet(doc, "Unit tests chạy nhanh → debug dễ khi thay đổi công thức")
    add_bullet(doc, "RAG eval đo chất lượng đầu cuối → phát hiện degradation sau khi update model/data")
    add_bullet(doc, "Robustness dataset → kiểm tra hành vi ngoài domain trước khi deploy")

    # Save
    output = os.path.abspath(OUTPUT_PATH)
    doc.save(output)
    print(f"Saved: {output}")


if __name__ == "__main__":
    build_doc()
