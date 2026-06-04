"""Generate simple Word report organized by Auepora framework (2405.07437v2).

Structure:
  1. Retrieval: Relevance (Context Precision/Recall) + Accuracy (MAP, MRR, Hit@K)
  2. Generation: Relevance (Answer Relevancy) + Faithfulness + Correctness
  3. Additional Requirements: Latency + Noise Robustness + Negative Rejection
"""

import json
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

RESULTS_DIR = Path("tests/evaluation/results")
OUTPUT_PATH = RESULTS_DIR / "bao-cao-auepora-don-gian.docx"

S1 = RESULTS_DIR / "eval-summary-20260602_173030.json"   # Q1-50
S2 = RESULTS_DIR / "eval-summary-20260602_193653.json"   # Q51-100
SR = RESULTS_DIR / "eval-summary-20260604_173359.json"   # Robustness


def shd(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    s = OxmlElement("w:shd")
    s.set(qn("w:val"), "clear")
    s.set(qn("w:color"), "auto")
    s.set(qn("w:fill"), hex_color)
    tcPr.append(s)


def border(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcB = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "4")
        b.set(qn("w:color"), "AAAAAA")
        tcB.append(b)
    tcPr.append(tcB)


def color(score):
    if score >= 0.80: return "C6EFCE"
    if score >= 0.60: return "FFEB9C"
    return "FFC7CE"


def avg(s1, s2, key):
    return round((s1.get(key, 0) + s2.get(key, 0)) / 2, 4)


def header_row(table, cols, bg="1F4E79"):
    hdr = table.rows[0].cells
    for cell, text in zip(hdr, cols):
        cell.text = text
        run = cell.paragraphs[0].runs[0]
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        shd(cell, bg)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER


def score_row(table, label, scores, note=""):
    row = table.add_row().cells
    row[0].text = label
    border(row[0])
    for i, sc in enumerate(scores, 1):
        row[i].text = f"{sc:.3f}" if isinstance(sc, float) else str(sc)
        row[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        if isinstance(sc, float):
            shd(row[i], color(sc))
        border(row[i])
    if note and len(row) > len(scores) + 1:
        row[-1].text = note
        border(row[-1])


def add_section_header(doc, text, color_hex="1F4E79"):
    p = doc.add_paragraph()
    run = p.add_run(f"  {text}  ")
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    pPr = p._p.get_or_add_pPr()
    shading = OxmlElement('w:shd')
    shading.set(qn('w:val'), 'clear')
    shading.set(qn('w:color'), 'auto')
    shading.set(qn('w:fill'), color_hex)
    pPr.append(shading)
    return p


def main():
    with open(S1) as f: s1 = json.load(f)
    with open(S2) as f: s2 = json.load(f)
    with open(SR) as f: sr = json.load(f)

    m1 = s1["metrics"]
    m2 = s2["metrics"]
    rr1 = s1.get("retrieval_rank_metrics", {})
    rr2 = s2.get("retrieval_rank_metrics", {})
    mr = sr["metrics"]

    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2)

    # Title
    t = doc.add_heading("KẾT QUẢ KIỂM THỬ HỆ THỐNG RAG CHATBOT", 0)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph("Theo Auepora Framework (2405.07437v2) — UIT CITD Chatbot")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].italic = True
    doc.add_paragraph()

    # Overview box
    info = doc.add_paragraph()
    info.add_run("Dataset: ").bold = True
    info.add_run("100 câu chuẩn + 11 câu robustness    ")
    info.add_run("Model: ").bold = True
    info.add_run("gemini-2.5-flash (judge) + BAAI/bge-m3 (embed) + bge-reranker-v2-m3    ")
    info.add_run("Ngày: ").bold = True
    info.add_run("02–04/06/2026")

    doc.add_paragraph()

    # ─── BLOCK 1: RETRIEVAL ───────────────────────────────────────────
    add_section_header(doc, "I. RETRIEVAL COMPONENT", "1F4E79")

    doc.add_paragraph()
    doc.add_paragraph("1.1  Retrieval Relevance — đánh giá bằng LLM-as-Judge").runs[0].bold = True
    p = doc.add_paragraph("Đo mức độ liên quan của các đoạn văn bản được truy xuất so với câu hỏi.")
    p.runs[0].italic = True

    t1 = doc.add_table(rows=1, cols=4)
    t1.style = "Table Grid"
    header_row(t1, ["Metric", "Batch 1 (Q1–50)", "Batch 2 (Q51–100)", "Trung bình"])
    score_row(t1, "Context Precision", [m1["context_precision"], m2["context_precision"],
              avg(m1, m2, "context_precision")])
    score_row(t1, "Context Recall", [m1["context_recall"], m2["context_recall"],
              avg(m1, m2, "context_recall")])
    doc.add_paragraph()

    doc.add_paragraph("1.2  Retrieval Accuracy — rank-based metrics (MAP, MRR, Hit@K)").runs[0].bold = True
    p = doc.add_paragraph("Đo độ chính xác vị trí xếp hạng của các chunk liên quan trong kết quả truy xuất.")
    p.runs[0].italic = True

    t2 = doc.add_table(rows=1, cols=4)
    t2.style = "Table Grid"
    header_row(t2, ["Metric", "Batch 1 (Q1–50)", "Batch 2 (Q51–100)", "Trung bình"])
    score_row(t2, "MAP@10", [rr1.get("map_at_10", 0), rr2.get("map_at_10", 0),
              round((rr1.get("map_at_10", 0) + rr2.get("map_at_10", 0)) / 2, 3)])
    score_row(t2, "MRR", [rr1.get("mrr", 0), rr2.get("mrr", 0),
              round((rr1.get("mrr", 0) + rr2.get("mrr", 0)) / 2, 3)])
    score_row(t2, "Hit@10", [rr1.get("hit_at_10", 0), rr2.get("hit_at_10", 0),
              round((rr1.get("hit_at_10", 0) + rr2.get("hit_at_10", 0)) / 2, 3)])
    doc.add_paragraph()

    # ─── BLOCK 2: GENERATION ─────────────────────────────────────────
    add_section_header(doc, "II. GENERATION COMPONENT", "375623")

    doc.add_paragraph()
    doc.add_paragraph("Đánh giá chất lượng câu trả lời của LLM dựa trên 3 chiều:").runs[0].bold = True

    t3 = doc.add_table(rows=1, cols=5)
    t3.style = "Table Grid"
    header_row(t3, ["Metric", "Ý nghĩa (Auepora)", "Batch 1", "Batch 2", "TB"], "375623")

    gen_rows = [
        ("Answer Relevancy", "Response ↔ Query", "answer_relevancy"),
        ("Faithfulness", "Response ↔ Relevant Docs", "faithfulness"),
        ("Correctness", "Response ↔ Ground Truth", "correctness"),
    ]
    for label, meaning, key in gen_rows:
        row = t3.add_row().cells
        row[0].text = label
        row[0].paragraphs[0].runs[0].bold = True
        border(row[0])
        row[1].text = meaning
        border(row[1])
        v1, v2 = m1.get(key, 0), m2.get(key, 0)
        va = round((v1 + v2) / 2, 3)
        for i, sc in enumerate([v1, v2, va], 2):
            row[i].text = f"{sc:.3f}"
            row[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            shd(row[i], color(sc))
            border(row[i])

    doc.add_paragraph()

    # ─── BLOCK 3: ADDITIONAL REQUIREMENTS ───────────────────────────
    add_section_header(doc, "III. ADDITIONAL REQUIREMENTS", "7030A0")

    doc.add_paragraph()
    doc.add_paragraph("3.1  Robustness — 11 câu (5 noise + 6 negative out-of-domain)").runs[0].bold = True

    t4 = doc.add_table(rows=1, cols=3)
    t4.style = "Table Grid"
    header_row(t4, ["Metric", "Điểm", "Ý nghĩa"], "7030A0")
    rob_rows = [
        ("Noise Robustness", mr.get("noise_robustness", 0), "Câu trả lời đúng khi context bị nhiễu"),
        ("Negative Rejection", mr.get("negative_rejection", 0), "Từ chối 100% câu ngoài domain"),
    ]
    for label, sc, meaning in rob_rows:
        row = t4.add_row().cells
        row[0].text = label
        row[0].paragraphs[0].runs[0].bold = True
        border(row[0])
        row[1].text = f"{sc:.3f}"
        row[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        shd(row[1], color(sc))
        border(row[1])
        row[2].text = meaning
        border(row[2])

    doc.add_paragraph()
    doc.add_paragraph("3.2  Latency").runs[0].bold = True

    t5 = doc.add_table(rows=1, cols=3)
    t5.style = "Table Grid"
    header_row(t5, ["", "Batch 1 (Q1–50)", "Batch 2 (Q51–100)"], "7030A0")
    lat1 = s1.get("latency", {})
    lat2 = s2.get("latency", {})
    for key, label in [("avg_ms", "Trung bình"), ("p50_ms", "P50"), ("p90_ms", "P90")]:
        row = t5.add_row().cells
        row[0].text = label
        border(row[0])
        for i, lat in enumerate([lat1, lat2], 1):
            val = lat.get(key, 0) / 1000
            row[i].text = f"{val:.1f}s"
            row[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            border(row[i])

    doc.add_paragraph()

    # ─── BLOCK 4: SUMMARY TABLE ──────────────────────────────────────
    add_section_header(doc, "IV. TỔNG KẾT", "C55A11")

    doc.add_paragraph()
    summary_rows = [
        # (metric, batch1, batch2, avg, status)
        ("Context Precision", m1["context_precision"], m2["context_precision"], None, "Retrieval"),
        ("Context Recall", m1["context_recall"], m2["context_recall"], None, "Retrieval"),
        ("MAP@10", rr1.get("map_at_10", 0), rr2.get("map_at_10", 0), None, "Retrieval"),
        ("MRR", rr1.get("mrr", 0), rr2.get("mrr", 0), None, "Retrieval"),
        ("Hit@10", rr1.get("hit_at_10", 0), rr2.get("hit_at_10", 0), None, "Retrieval"),
        ("Answer Relevancy", m1["answer_relevancy"], m2["answer_relevancy"], None, "Generation"),
        ("Faithfulness", m1["faithfulness"], m2["faithfulness"], None, "Generation"),
        ("Correctness", m1["correctness"], m2["correctness"], None, "Generation"),
        ("Noise Robustness", mr.get("noise_robustness", 0), None, None, "Robustness"),
        ("Negative Rejection", mr.get("negative_rejection", 0), None, None, "Robustness"),
    ]

    t6 = doc.add_table(rows=1, cols=5)
    t6.style = "Table Grid"
    header_row(t6, ["Nhóm", "Metric", "Batch 1", "Batch 2", "Đánh giá"], "C55A11")

    for metric, v1, v2, _, group in summary_rows:
        row = t6.add_row().cells
        row[0].text = group
        border(row[0])
        row[1].text = metric
        border(row[1])
        row[2].text = f"{v1:.3f}"
        row[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        shd(row[2], color(v1))
        border(row[2])
        if v2 is not None:
            row[3].text = f"{v2:.3f}"
            row[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            shd(row[3], color(v2))
        else:
            row[3].text = "N/A"
            row[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        border(row[3])
        sc = v1 if v2 is None else (v1 + v2) / 2
        row[4].text = "✅ Tốt" if sc >= 0.75 else ("⚠️ Khá" if sc >= 0.55 else "❌ Yếu")
        row[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        border(row[4])

    doc.add_paragraph()
    foot = doc.add_paragraph("Tham chiếu: Yu et al. (2024) — Evaluation of RAG: A Survey (arXiv:2405.07437v2)")
    foot.runs[0].italic = True
    foot.runs[0].font.size = Pt(9)
    foot.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.save(OUTPUT_PATH)
    print(f"Saved: {OUTPUT_PATH}")


if __name__ == "__main__":
    main()
