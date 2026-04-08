"""Extract text from DOCX and JSON files in data/raw, merge into all_documents_ocr.json.

Usage: python3 data/raw/extract-docx-json.py
"""

import json
import re
from pathlib import Path

from docx import Document

RAW_DIR = Path(__file__).parent
OUTPUT_DIR = RAW_DIR.parent / "processed"
OCR_FILE = OUTPUT_DIR / "all_documents_ocr.json"


def clean_text(text: str) -> str:
    """Remove noise, collapse whitespace, drop short lines."""
    text = re.sub(
        r"[^\w\sàáảãạăắằẳẵặâấầẩẫậèéẻẽẹêếềểễệìíỉĩịòóỏõọôốồổỗộơớờởỡợùúủũụưứừửữựỳýỷỹỵđÀÁẢÃẠĂẮẰẲẴẶÂẤẦẨẪẬÈÉẺẼẸÊẾỀỂỄỆÌÍỈĨỊÒÓỎÕỌÔỐỒỔỖỘƠỚỜỞỠỢÙÚỦŨỤƯỨỪỬỮỰỲÝỶỸỴĐ.,;:!?()\"'\-/§°%&@#\d]",
        "",
        text,
    )
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    lines = [l for l in text.split("\n") if len(l.strip()) > 2]
    return "\n".join(lines).strip()


def extract_docx(docx_path: Path) -> list[dict]:
    """Extract text from DOCX, chunk by ~2000 chars per paragraph group."""
    doc = Document(str(docx_path))
    # Collect all paragraphs with text
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]

    # Also extract text from tables
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                paragraphs.append(row_text)

    if not paragraphs:
        return []

    # Group paragraphs into chunks of ~2000 chars
    chunks = []
    current_chunk = []
    current_len = 0
    chunk_num = 1

    for para in paragraphs:
        if current_len + len(para) > 2000 and current_chunk:
            text = clean_text("\n".join(current_chunk))
            if text:
                chunks.append({
                    "content": text,
                    "page": chunk_num,
                    "source": docx_path.name,
                    "method": "python_docx",
                })
                chunk_num += 1
            current_chunk = [para]
            current_len = len(para)
        else:
            current_chunk.append(para)
            current_len += len(para)

    # Last chunk
    if current_chunk:
        text = clean_text("\n".join(current_chunk))
        if text:
            chunks.append({
                "content": text,
                "page": chunk_num,
                "source": docx_path.name,
                "method": "python_docx",
            })

    return chunks


def extract_json(json_path: Path) -> list[dict]:
    """Extract text from JSON array with {url, title, text} structure."""
    with open(json_path, encoding="utf-8") as f:
        data = json.load(f)

    chunks = []
    for i, item in enumerate(data, 1):
        title = item.get("title", "")
        text = item.get("text", "")
        combined = f"{title}\n\n{text}" if title else text
        combined = clean_text(combined)

        if combined:
            chunks.append({
                "content": combined,
                "page": i,
                "source": json_path.name,
                "method": "json_extract",
            })

    return chunks


def main():
    # Load existing OCR data
    existing = []
    if OCR_FILE.exists():
        with open(OCR_FILE, encoding="utf-8") as f:
            existing = json.load(f)
        print(f"Loaded {len(existing)} existing chunks from OCR")

    new_chunks = []

    # Process DOCX files
    docx_files = sorted(RAW_DIR.glob("*.docx"))
    print(f"\nFound {len(docx_files)} DOCX files")
    for docx_path in docx_files:
        chunks = extract_docx(docx_path)
        print(f"  {docx_path.name}: {len(chunks)} chunks")
        new_chunks.extend(chunks)

    # Process JSON files
    json_files = sorted(RAW_DIR.glob("*.json"))
    print(f"\nFound {len(json_files)} JSON files")
    for json_path in json_files:
        chunks = extract_json(json_path)
        print(f"  {json_path.name}: {len(chunks)} chunks")
        new_chunks.extend(chunks)

    # Merge and save
    all_chunks = existing + new_chunks
    with open(OCR_FILE, "w", encoding="utf-8") as f:
        json.dump(all_chunks, f, ensure_ascii=False, indent=2)

    print(f"\n{'=' * 60}")
    print(f"Added {len(new_chunks)} new chunks")
    print(f"Total: {len(all_chunks)} chunks in {OCR_FILE}")


if __name__ == "__main__":
    main()
